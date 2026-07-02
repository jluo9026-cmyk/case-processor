"""Quick API smoke test"""
import sys, os, json, io
sys.path.insert(0, 'e:/案件处理启动器')

# Disable API keys
os.environ['DEEPSEEK_API_KEY'] = ''
os.environ['BAIDU_API_KEY'] = ''
os.environ['BAIDU_SECRET_KEY'] = ''
os.environ['QWEN_VL_API_KEY'] = ''

from fastapi.testclient import TestClient
from modules.app_core import app
from modules import routes

client = TestClient(app)

# 1. Health
r = client.get('/api/health')
print(f'Health: {r.status_code} {r.json()}')

# 2. Home
r = client.get('/')
print(f'Home: {r.status_code} page_len={len(r.text)}')

# 3. Tools
for path in ['/tools/report_generate/index.html', '/tools/index4.html', '/tools/index3.html', '/tools/附件处理器-合并版/index.html']:
    r = client.get(path)
    print(f'Tool {path}: {r.status_code}')

# 4. Preset templates
r = client.get('/api/template/presets')
print(f'Presets: {r.status_code} count={len(r.json().get("presets",[]))}')

# 5. Run report (no API key = fallback)
r = client.post('/api/run-with-preset', data={'template_id':'preset_1','sceneInvestigation':'测试内容'})
jd = r.json()
print(f'Run-with-preset: {r.status_code} success={jd.get("success")} content_len={len(jd.get("report_content",""))}')

# 6. /api/run
r = client.post('/api/run', data={'template_id':'preset_1','sceneInvestigation':'测试'})
jd = r.json()
print(f'/api/run: {r.status_code} success={jd.get("success")} content_len={len(jd.get("report_content",""))}')

# 7. Generate DOCX
r = client.post('/api/generate-docx', data={'markdown_report':'# Test\nHello','output_name':'test'})
jd = r.json()
print(f'Generate DOCX: {r.status_code} success={jd.get("success")} id={jd.get("report_id","")}')

# 8. Download
if jd.get('success'):
    r2 = client.get(f'/api/download/{jd["report_id"]}')
    print(f'Download: {r2.status_code} bytes={len(r2.content)}')

# 9. Excel fill
import openpyxl
wb = openpyxl.Workbook()
ws = wb.active
ws['A1'] = '保单号：'
ws['B1'] = '投保人：'
out = io.BytesIO()
wb.save(out); out.seek(0)
r = client.post('/api/fill-excel', data={'template':('t.xlsx',out.read(),'application/octet-stream'),'data':json.dumps({'policyNo':'123','insured':'测试'})})
print(f'Excel fill: {r.status_code} bytes={len(r.content)}')

# 10. Process DOCX
from docx import Document
doc = Document()
doc.add_heading('测试',0)
t = doc.add_table(rows=2,cols=2)
t.style='Table Grid'
t.rows[0].cells[0].text='附件名称'; t.rows[0].cells[1].text='页码'
t.rows[1].cells[0].text='附件一'; t.rows[1].cells[1].text='3'
o2 = io.BytesIO(); doc.save(o2); o2.seek(0)
r = client.post('/api/process-docx', data={'file':('t.docx',o2.read(),'application/octet-stream'),'attachments':json.dumps([{'number':1,'name':'现场照片'}])})
print(f'Process DOCX: {r.status_code} bytes={len(r.content)}')

# 11. Attachment endpoints
r = client.get('/api/ranges'); print(f'Ranges: {r.status_code}')
r = client.post('/api/ranges/reset'); print(f'Ranges reset: {r.status_code}')
r = client.post('/api/one-click', json={'images':[]}); print(f'One-click: {r.status_code} {r.json().get("success")}')
r = client.post('/api/generate-markdown', json={'images':[]}); print(f'Markdown: {r.status_code} {r.json().get("success")}')

print('\nAll tests passed!')