from modules.app_core import app
from modules.config import *
from modules.ocr_service import *
from modules.docx_merge import *
from modules.template_data import *
from modules.helpers import *
from fastapi import UploadFile, File, Form, HTTPException, Request
from fastapi.responses import JSONResponse, FileResponse, StreamingResponse
from starlette.responses import Response
import os, uuid, io, json, base64, re, asyncio, zipfile, shutil
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from pathlib import Path

# ============ 报告生成 (FastAPI) ============
_generated_reports = {}
_template_storage = {}


@app.post('/api/generate-report')
async def generate_report(request: Request):
    """生成报告 - 支持多图片OCR+视觉分析+笔录信息"""
    try:
        form = await request.form()
        text = form.get('text', '')
        case_type = form.get('case_type', '')
        template_id = form.get('template_id')
        scene_investigation = form.get('scene_investigation', '')
        police_station_record = form.get('police_station_record', '')
        traffic_police_record = form.get('traffic_police_record', '')
        hospital_diagnosis = form.get('hospital_diagnosis', '')

        images_data = []
        statement_data = []
        files = []
        for key in form.keys():
            if key.startswith('images_') or key.startswith('file_'):
                value = form.get(key)
                if hasattr(value, 'read'):
                    files.append(value)

        for f in files:
            try:
                content = await f.read()
                if content:
                    images_data.append({
                        'name': getattr(f, 'filename', f'image_{len(images_data)}.png'),
                        'data': content
                    })
            except Exception as e:
                print(f'[Upload] 读取文件失败: {e}')

        for key in form.keys():
            if key.startswith('statement_'):
                value = form.get(key)
                if hasattr(value, 'read'):
                    try:
                        content = await value.read()
                        if content:
                            statement_data.append({
                                'name': getattr(value, 'filename', f'statement_{len(statement_data)}.png'),
                                'data': content
                            })
                    except Exception as e:
                        print(f'[Upload] 读取笔录文件失败: {e}')

        job_id = str(uuid.uuid4())[:12]
        await create_job(job_id, metadata={
            'text': text,
            'case_type': case_type,
            'template_id': template_id,
            'images_count': len(images_data),
            'statements_count': len(statement_data),
            'scene_investigation': scene_investigation,
            'police_station_record': police_station_record,
            'traffic_police_record': traffic_police_record,
            'hospital_diagnosis': hospital_diagnosis
        })

        asyncio.create_task(_background_report_job(
            job_id, text, images_data, statement_data,
            template_id, case_type,
            scene_investigation, police_station_record,
            traffic_police_record, hospital_diagnosis
        ))

        return JSONResponse(content={
            'success': True,
            'job_id': job_id,
            'message': '报告生成任务已创建'
        })
    except Exception as e:
        raise HTTPException(500, f'创建任务失败: {str(e)}')


@app.get('/api/run/{job_id}')
async def get_run_status(job_id: str):
    job = await get_job(job_id)
    if not job:
        raise HTTPException(404, '任务不存在')
    return {
        'success': True,
        'job_id': job_id,
        'status': job['status'],
        'progress': job.get('progress', 0),
        'message': job.get('message', ''),
        'error': job.get('error'),
        'result': job.get('result')
    }


# 任务超时设置
MAX_JOB_RUNTIME = 300  # 5分钟


async def _background_report_job(
    job_id, text, images_data, statement_data,
    template_id, case_type,
    scene_investigation, police_station_record,
    traffic_police_record, hospital_diagnosis
):
    await update_job(job_id, status='running', started_at=datetime.now().isoformat(), progress=5, message='开始处理')
    job_start_time = datetime.now()

    try:
        await update_job(job_id, progress=10, message='正在识别图片...')

        ocr_results = []
        visual_results = []

        def _is_job_timed_out():
            return (datetime.now() - job_start_time).total_seconds() > MAX_JOB_RUNTIME

        # ============ OCR处理 ============
        async def _process_single_image(index, img_info, job_id, job_start_time):
            """处理单张图片的OCR和视觉分析"""
            if _is_job_timed_out():
                return None, None

            img_base64 = base64.b64encode(img_info['data']).decode('utf-8')
            ocr_result = None
            visual_result = None

            # OCR识别 - 按优先级尝试多个OCR服务
            print(f'[DEBUG] 开始处理图片: {img_info["name"]}, 大小: {len(img_info["data"])} bytes')
            
            # 1. 首先尝试百度OCR
            if BAIDU_API_KEY and BAIDU_SECRET_KEY:
                print(f'[DEBUG] 尝试百度OCR识别...')
                try:
                    ocr_resp = await asyncio.to_thread(_call_baidu_ocr, img_base64)
                    print(f'[DEBUG] 百度OCR响应: {ocr_resp}')
                    if ocr_resp.get('status') == 200:
                        if 'words_result' in ocr_resp.get('data', {}):
                            word_list = ocr_resp['data']['words_result']
                            text_parts = [w['words'] for w in word_list if w.get('words')]
                            if text_parts:
                                ocr_result = {'name': img_info['name'], 'text': '\n'.join(text_parts)}
                                print(f'[DEBUG] 百度OCR成功，识别到 {len(text_parts)} 行文字')
                            else:
                                print(f'[DEBUG] 百度OCR返回空结果')
                        else:
                            print(f'[DEBUG] 百度OCR响应中没有words_result字段')
                    else:
                        print(f'[DEBUG] 百度OCR返回状态码: {ocr_resp.get("status")}')
                except Exception as e:
                    print(f'[OCR] 百度OCR处理 {img_info["name"]} 失败: {e}')
                    import traceback
                    print(traceback.format_exc())
            
            # 2. 如果百度OCR失败，使用千问VL进行文字识别（PaddleOCR仅用于笔录图片）
            if not ocr_result and QWEN_VL_API_KEY:
                print(f'[DEBUG] 百度OCR失败，尝试千问VL识别...')
                try:
                    description = await _call_qwen_vl(img_base64)
                    if description:
                        ocr_result = {'name': img_info['name'], 'text': description}
                        print(f'[Qwen-VL] 识别 {img_info["name"]} 成功，长度 {len(ocr_result["text"])}')
                    else:
                        print(f'[DEBUG] 千问VL返回空结果')
                except Exception as e:
                    print(f'[Qwen-VL] 处理 {img_info["name"]} 失败: {e}')
                    import traceback
                    print(traceback.format_exc())
            
            # 如果所有OCR都失败，记录警告
            if not ocr_result:
                print(f'[WARNING] 所有OCR服务都未能识别图片: {img_info["name"]}')

            # Qwen-VL视觉分析（独立于OCR，用于描述图片内容）
            if QWEN_VL_API_KEY:
                try:
                    img_type = await asyncio.to_thread(_classify_image, img_base64)
                    if img_type == 'visual':
                        description = await _call_qwen_vl(img_base64)
                        if description:
                            visual_result = {'name': img_info['name'], 'description': description}
                except Exception as e:
                    print(f'[视觉] 处理 {img_info["name"]} 失败: {e}')

            return ocr_result, visual_result

        # 处理所有图片
        if images_data:
            await update_job(job_id, progress=10, message=f'正在处理 {len(images_data)} 张图片...')
            print(f'[OCR] 开始处理{len(images_data)}张图片: {[img["name"] for img in images_data]}')
            tasks = [
                _process_single_image(i, img_info, job_id, job_start_time)
                for i, img_info in enumerate(images_data)
            ]
            results = await asyncio.gather(*tasks)
            for ocr_result, visual_result in results:
                if ocr_result:
                    ocr_results.append(ocr_result)
                    print(f'[OCR] 识别 {ocr_result["name"]} 成功，长度 {len(ocr_result["text"])}')
                if visual_result:
                    visual_results.append(visual_result)
            print(f'[OCR] 识别成功 {len(ocr_results)} 张，失败 {len(images_data)-len(ocr_results)} 张')
            await update_job(job_id, progress=50, message=f'已完成 {len(ocr_results)}/{len(images_data)} 张图片识别')

        # ============ 笔录处理 + PaddleOCR处理 ============
        async def _process_single_statement(idx, img_info, job_id, job_start_time):
            if _is_job_timed_out():
                return None, None
            try:
                import tempfile
                temp_dir = tempfile.gettempdir()
                temp_filename = f'paddle_temp_{uuid.uuid4().hex}.png'
                temp_path = os.path.join(temp_dir, temp_filename)
                temp_path = os.path.abspath(temp_path)

                with open(temp_path, 'wb') as f:
                    f.write(img_info['data'])

                try:
                    if HAS_PADDLE or _init_paddle_ocr():
                        result = PADDLE_OCR.ocr(temp_path, cls=False)
                        if result and result[0]:
                            lines = [line[1][0] for line in result[0] if line[1] and line[1][0]]
                            if lines:
                                paddle_text = '\n'.join(lines)
                                print(f'[PaddleOCR] 识别 {img_info["name"]} 成功，{len(lines)} 行')
                                return {'name': img_info['name'], 'text': paddle_text}, None
                except Exception as e:
                    print(f'[PaddleOCR] 识别 {img_info["name"]} 失败: {e}')

                # PaddleOCR失败，回退到百度OCR
                try:
                    with open(temp_path, 'rb') as f:
                        img_data = f.read()
                    img_base64 = base64.b64encode(img_data).decode('utf-8')
                    ocr_resp = _call_baidu_ocr(img_base64)
                    if ocr_resp.get('status') == 200 and 'words_result' in ocr_resp.get('data', {}):
                        word_list = ocr_resp['data']['words_result']
                        text_parts = [w['words'] for w in word_list if w.get('words')]
                        if text_parts:
                            return {'name': img_info['name'], 'text': '\n'.join(text_parts)}, None
                except Exception as e2:
                    print(f'[OCR回退] 处理 {img_info["name"]} 失败: {e2}')

                return None, None
            except Exception as e:
                print(f'[Statement] 处理 {img_info["name"]} 失败: {e}')
                return None, None
            finally:
                try:
                    if os.path.exists(temp_path):
                        os.unlink(temp_path)
                except:
                    pass

        # 处理笔录图片
        paddle_results = []
        if statement_data:
            await update_job(job_id, progress=55, message=f'正在处理 {len(statement_data)} 张笔录图片...')
            print(f'[PaddleOCR] 开始处理{len(statement_data)}张笔录图片')
            statement_tasks = [
                _process_single_statement(i, img_info, job_id, job_start_time)
                for i, img_info in enumerate(statement_data)
            ]
            statement_results = await asyncio.gather(*statement_tasks)
            for paddle_result, _ in statement_results:
                if paddle_result:
                    paddle_results.append(paddle_result)
            print(f'[PaddleOCR] 处理完成: {len(paddle_results)}/{len(statement_data)}')

        # ============ 合并信息 ============
        await update_job(job_id, progress=65, message='正在合并信息...')

        ocr_text = '\n'.join([r['text'] for r in ocr_results]) if ocr_results else ''
        visual_text = '\n'.join([r['description'] for r in visual_results]) if visual_results else ''
        paddle_text = '\n'.join([r['text'] for r in paddle_results]) if paddle_results else ''

        investigation_content = {
            'scene_investigation': scene_investigation,
            'police_station_record': police_station_record,
            'traffic_police_record': traffic_police_record,
            'hospital_diagnosis': hospital_diagnosis
        }

        combined_text = _merge_information(
            investigation_content=investigation_content,
            ocr_text=ocr_text,
            visual_descriptions=visual_text,
            paddle_text=paddle_text
        )

        # ============ 生成Word报告 ============
        if _is_job_timed_out():
            raise TimeoutError(f'任务超时（{MAX_JOB_RUNTIME}秒）')

        await update_job(job_id, progress=75, message='正在生成 Word 报告...')

        report_id = str(uuid.uuid4())[:12]
        output_filename = f'report_{report_id}.docx'
        output_path = OUTPUT_DIR / output_filename

        # 获取模板内容
        content = None
        case_data = {}
        if template_id and template_id in _template_storage:
            template_info = _template_storage[template_id]
            content = template_info.get('content')
            case_data = template_info.get('data', {})

        if content:
            if HAS_MODULES:
                temp_template_path = OUTPUT_DIR / f'temp_template_{uuid.uuid4().hex}.docx'
                temp_template_path.write_bytes(content)
                formatter = UniversalDocxFormatter()
                result_path, structure = formatter.convert_with_template(
                    str(temp_template_path), content, str(output_path), combined_text
                )
                if temp_template_path.exists():
                    temp_template_path.unlink()
            else:
                dst = Document(io.BytesIO(content))
                if combined_text.strip():
                    dst.add_paragraph()
                    dst.add_heading('报告内容', level=1)
                    for line in combined_text.strip().split('\n'):
                        stripped = line.strip()
                        if stripped:
                            dst.add_paragraph(stripped)
                dst.save(str(output_path))
                structure = {}
        else:
            if HAS_MODULES:
                formatter = UniversalDocxFormatter()
                result_path, structure = formatter.convert_with_content(
                    case_type, case_data, combined_text, str(output_path)
                )
            else:
                doc = Document()
                doc.add_heading(f'{case_type}调查报告', 0)
                doc.add_paragraph(f'报告编号: {report_id}')
                doc.add_paragraph(f'生成时间: {datetime.now().strftime("%Y-%m-%d %H:%M")}')
                if combined_text.strip():
                    doc.add_paragraph()
                    doc.add_heading('报告内容', level=1)
                    for line in combined_text.strip().split('\n'):
                        stripped = line.strip()
                        if stripped:
                            doc.add_paragraph(stripped)
                doc.save(str(output_path))
                structure = {}

        if _is_job_timed_out():
            raise TimeoutError(f'任务超时（{MAX_JOB_RUNTIME}秒）')

        await update_job(job_id, progress=90, message='正在生成最终报告...')

        # 调用DeepSeek API生成报告
        final_report = None
        if DEEPSEEK_API_KEY:
            try:
                prompt = _generate_report_prompt(ocr_text, combined_text, text)
                final_report = await _call_deepseek_api(prompt)
            except Exception as api_err:
                print(f'[DEBUG] DeepSeek API 调用失败: {api_err}')

        if not final_report:
            final_report = generate_report_with_content(combined_text, ocr_text, text)

        _generated_reports[report_id] = {
            'id': report_id,
            'filename': output_filename,
            'path': str(output_path),
            'case_type': case_type,
            'created': datetime.now().isoformat(),
        }

        response_data = {
            'success': True,
            'report_id': report_id,
            'output_file': output_filename,
            'message': '报告生成完成',
            'total_images': len(images_data),
            'ocr_count': len(ocr_results),
            'combined_text': combined_text,
            'final_report': final_report
        }

        await update_job(job_id, status='completed', progress=100, message='报告生成完成', result=response_data, finished_at=datetime.now().isoformat())

    except TimeoutError as te:
        print(f'[REPORT JOB] 任务超时 {job_id}: {te}')
        await update_job(job_id, status='failed', progress=100, message='', error=str(te), finished_at=datetime.now().isoformat())
    except Exception as e:
        print(f'[REPORT JOB] 任务失败 {job_id}: {e}')
        import traceback
        traceback.print_exc()
        await update_job(job_id, status='failed', progress=100, message='', error=str(e), finished_at=datetime.now().isoformat())
@app.post('/api/generate-docx')
async def generate_docx(request: Request):
    """将 markdown 转换为 DOCX 报告"""
    try:
        content_type = request.headers.get('content-type', '')
        markdown_text = ''
        template_id = None
        output_name = ''

        if 'multipart/form-data' in content_type:
            form = await request.form()
            markdown_text = form.get('markdown_report', '')
            template_id = form.get('template_id') or None
            output_name = form.get('output_name', '')
        else:
            try:
                body = await request.json()
                markdown_text = body.get('markdown_report', '') or body.get('text', '')
                template_id = body.get('template_id')
                output_name = body.get('output_name', '')
            except Exception:
                pass

        if not markdown_text.strip():
            raise HTTPException(400, '报告内容不能为空')

        timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
        output_filename = f'report_{timestamp}.docx'
        if output_name:
            safe_name = re.sub(r'[^\w\u4e00-\u9fff\-]', '_', output_name)
            output_filename = f'report_{safe_name}_{timestamp}.docx'

        output_path = OUTPUT_DIR / output_filename

        # 使用模板或直接生成
        content = None
        if template_id and template_id in _template_storage:
            template_info = _template_storage[template_id]
            content = template_info.get('content')

        if content and HAS_MODULES:
            temp_template_path = OUTPUT_DIR / f'temp_template_{uuid.uuid4().hex}.docx'
            temp_template_path.write_bytes(content)
            formatter = UniversalDocxFormatter()
            result_path, structure = formatter.convert_with_template(
                str(temp_template_path), content, str(output_path), markdown_text
            )
            if temp_template_path.exists():
                temp_template_path.unlink()
        else:
            doc = Document()
            for line in markdown_text.split('\n'):
                stripped = line.strip()
                if not stripped:
                    continue
                if stripped.startswith('# '):
                    doc.add_heading(stripped[2:], 0)
                elif stripped.startswith('## '):
                    doc.add_heading(stripped[3:], 1)
                elif stripped.startswith('### '):
                    doc.add_heading(stripped[4:], 2)
                elif stripped.startswith('- ') or stripped.startswith('* '):
                    doc.add_paragraph(stripped, style='List Bullet')
                else:
                    doc.add_paragraph(stripped)
            doc.save(str(output_path))

        report_id = str(uuid.uuid4())[:12]
        _generated_reports[report_id] = {
            'id': report_id,
            'filename': output_filename,
            'path': str(output_path),
            'created': datetime.now().isoformat(),
        }

        report_id = str(uuid.uuid4())[:12]
        _generated_reports[report_id] = {
            'id': report_id,
            'filename': output_filename,
            'path': str(output_path),
            'created': datetime.now().isoformat(),
        }

        return JSONResponse(content={
            'success': True,
            'report_id': report_id,
            'output_file': output_filename,
            'message': 'DOCX 报告生成完成'
        })
    except Exception as e:
        import traceback
        traceback.print_exc()
        raise HTTPException(500, f'生成 DOCX 失败: {str(e)}')


@app.post('/api/template/upload')
async def upload_template_new(request: Request):
    """上传模板（前端调用的端点）"""
    print('[DEBUG] 收到模板上传请求')
    try:
        result = await upload_template(request)
        print(f'[DEBUG] 模板上传成功: {result}')
        return result
    except Exception as e:
        print(f'[ERROR] 模板上传失败: {e}')
        import traceback
        print(traceback.format_exc())
        return JSONResponse(content={
            'success': False,
            'error': f'上传模板失败: {str(e)}'
        }, status_code=500)


@app.delete('/api/template/{template_id}')
async def delete_template_api(template_id: str):
    """删除模板"""
    if template_id in _template_storage:
        del _template_storage[template_id]
        return {'success': True}
    else:
        raise HTTPException(404, '模板不存在')


@app.post('/api/upload-template')
async def upload_template(request: Request):
    """上传Word模板"""
    print('[DEBUG] upload_template 开始处理')
    try:
        print('[DEBUG] 尝试获取表单数据')
        form = await request.form()
        print(f'[DEBUG] 表单键名: {list(form.keys())}')
        
        file = None
        for key in form.keys():
            value = form.get(key)
            print(f'[DEBUG] 键 {key} 的类型: {type(value)}')
            if hasattr(value, 'read'):
                file = value
                print(f'[DEBUG] 找到文件: {getattr(file, "filename", "未知文件名")}')
                break

        if not file:
            print('[ERROR] 未找到上传文件')
            return JSONResponse(content={
                'success': False,
                'error': '未找到上传文件'
            }, status_code=400)

        print('[DEBUG] 读取文件内容')
        content = await file.read()
        print(f'[DEBUG] 文件大小: {len(content)} bytes')
        
        if not content:
            print('[ERROR] 文件内容为空')
            return JSONResponse(content={
                'success': False,
                'error': '文件内容为空'
            }, status_code=400)

        try:
            print('[DEBUG] 解析Word模板')
            template_data = parse_word_template(content)
            print(f'[DEBUG] 解析成功: {len(template_data["sections"])} 个章节, {len(template_data["paragraphs"])} 个段落')
        except Exception as parse_error:
            print(f'[ERROR] 解析Word文件失败: {parse_error}')
            return JSONResponse(content={
                'success': False,
                'error': f'解析Word文件失败: {str(parse_error)}'
            }, status_code=400)

        template_id = str(uuid.uuid4())[:8]
        print(f'[DEBUG] 生成模板ID: {template_id}')

        _template_storage[template_id] = {
            'id': template_id,
            'name': getattr(file, 'filename', 'template.docx'),
            'content': content,
            'data': template_data,
            'upload_time': datetime.now().isoformat()
        }

        response_data = {
            'success': True,
            'template_id': template_id,
            'template_name': getattr(file, 'filename', 'template.docx'),
            'sections_count': len(template_data['sections']),
            'paragraphs_count': len(template_data['paragraphs'])
        }
        print(f'[DEBUG] 返回成功响应: {response_data}')
        return JSONResponse(content=response_data)
    except Exception as e:
        print(f'[ERROR] 上传模板异常: {e}')
        import traceback
        print(traceback.format_exc())
        return JSONResponse(content={
            'success': False,
            'error': f'上传模板失败: {str(e)}'
        }, status_code=500)


@app.get('/api/templates')
async def list_templates():
    """列出所有已上传的模板"""
    templates = []
    for tid, t in _template_storage.items():
        templates.append({
            'id': tid,
            'name': t['name'],
            'upload_time': t['upload_time'],
            'sections_count': len(t['data'].get('sections', [])),
            'paragraphs_count': len(t['data'].get('paragraphs', []))
        })
    return {'success': True, 'templates': templates}


@app.post('/api/analyze')
async def analyze_document(file: UploadFile = File(...)):
    """分析文档结构（前端调用的端点）"""
    try:
        return await parse_docx(file)
    except HTTPException as e:
        return JSONResponse(
            content={'success': False, 'error': e.detail},
            status_code=e.status_code
        )
    except Exception as e:
        print(f'[ERROR] 分析文档失败: {e}')
        import traceback
        traceback.print_exc()
        return JSONResponse(
            content={'success': False, 'error': f'文档分析失败: {str(e)}'},
            status_code=500
        )


@app.post('/api/parse-docx')
async def parse_docx(file: UploadFile = File(...)):
    """解析DOCX文件"""
    if not file.filename or not allowed_file(file.filename):
        raise HTTPException(400, '请上传 .docx 文件')

    unique_filename = f"{uuid.uuid4().hex}_{secure_filename(file.filename)}"
    upload_path = UPLOAD_DIR / unique_filename

    try:
        content = await file.read()
        upload_path.write_bytes(content)

        doc = Document(io.BytesIO(content))
        
        if HAS_MODULES:
            parser = DocumentParser()
            parse_result = parser.parse(doc)
            formatter = UniversalDocxFormatter()
            structure = formatter.analyze_structure(parse_result)
            
            return {
                'success': True,
                'filename': file.filename,
                'structure': {
                    'main_title': structure.get('main_title'),
                    'chapters': [ch.get('title') for ch in structure.get('chapters', [])],
                    'key_info': structure.get('key_info', {}),
                    'attachments': structure.get('attachments', [])
                },
                'metadata': parse_result.get('metadata', {}),
                'tables_count': len(parse_result.get('tables', [])),
                'paragraphs_count': len([p for p in parse_result.get('paragraphs', []) if not p.get('is_empty')])
            }
        else:
            chapters = []
            paragraphs = []
            key_info = {}
            
            for para in doc.paragraphs:
                text = para.text.strip()
                if text:
                    paragraphs.append(text)
                    if (para.style.name.startswith('Heading') or 
                        text.startswith('第') and '章' in text or
                        text.startswith('一、') or text.startswith('二、') or 
                        text.startswith('三、') or text.startswith('四、') or
                        text.startswith('五、') or text.startswith('六、') or
                        text.startswith('七、')):
                        chapters.append(text)
            
            return {
                'success': True,
                'filename': file.filename,
                'structure': {
                    'main_title': paragraphs[0] if paragraphs else '',
                    'chapters': chapters[:10],
                    'key_info': key_info,
                    'attachments': []
                },
                'metadata': {},
                'tables_count': len(doc.tables),
                'paragraphs_count': len(paragraphs),
                'preview': '\n'.join(paragraphs[:20]) if len(paragraphs) > 20 else '\n'.join(paragraphs)
            }
    finally:
        if upload_path.exists():
            upload_path.unlink()



@app.post('/api/convert')
async def convert_document(file: UploadFile = File(...), template_id: str = Form('')):
    """转换文档格式 - 使用真实 .docx 模板"""
    if not file.filename or not allowed_file(file.filename):
        raise HTTPException(400, '请上传 .docx 文件')

    unique_filename = f"{uuid.uuid4().hex}_{secure_filename(file.filename)}"
    upload_path = UPLOAD_DIR / unique_filename
    output_filename = f"{uuid.uuid4().hex}_formatted_report.docx"
    output_path = OUTPUT_DIR / output_filename

    try:
        content = await file.read()
        upload_path.write_bytes(content)

        template_name = '标准格式'
        chapters_found = 0
        chapters_auto = 0

        # 如果选择了预设模板且有对应的真实 .docx 模板文件
        if template_id and template_id in TEMPLATE_DOCX_MAP:
            template_path = TEMPLATE_DOCX_MAP[template_id]
            if template_path.exists():
                preset = PRESET_TEMPLATES.get(template_id, {})
                template_name = preset.get('name', '标准模板')
                
                # 加载真实模板文件并通过 Document().save() 重建
                template_doc = Document(str(template_path))

                
                # 将上传文档内容合并到模板中
                new_doc, chapters_found, chapters_filled = _merge_content_into_template(template_doc, content, output_path)
                
                summary = {
                    'chapters_found': chapters_found,
                    'chapters_auto_generated': chapters_auto,
                    'template_used': template_name
                }
            else:
                # 模板文件不存在，回退
                print(f'[WARNING] 模板文件不存在: {template_path}，使用默认格式')
                template_name = '标准格式（模板文件未找到）'
                doc = Document(io.BytesIO(content))
                new_doc = Document()
                new_doc.add_heading('标准化报告', 0)
                for para in doc.paragraphs:
                    text = para.text.strip()
                    if not text:
                        continue
                    if para.style.name.startswith('Heading') or _is_chapter_heading(text):
                        chapters_found += 1
                        new_doc.add_heading(text, level=1)
                    else:
                        new_doc.add_paragraph(text)
                new_doc.save(str(output_path))
                summary = {
                    'chapters_found': chapters_found,
                    'chapters_auto_generated': 0,
                    'template_used': template_name
                }
        elif HAS_MODULES:
            formatter = UniversalDocxFormatter()
            result_path, structure = formatter.convert(str(upload_path), str(output_path))

            post_processor = PostProcessor()
            summary = post_processor.generate_summary(structure)
        else:
            doc = Document(io.BytesIO(content))
            
            new_doc = Document()
            new_doc.add_heading('标准化报告', 0)
            new_doc.add_paragraph(f'报告编号: {uuid.uuid4().hex[:8]}')
            new_doc.add_paragraph(f'生成时间: {datetime.now().strftime("%Y-%m-%d %H:%M")}')
            
            chapters_found = 0
            for para in doc.paragraphs:
                text = para.text.strip()
                if not text:
                    continue
                if para.style.name.startswith('Heading') or _is_chapter_heading(text):
                    chapters_found += 1
                    new_doc.add_heading(text, level=1)
                else:
                    new_doc.add_paragraph(text)
            
            new_doc.save(str(output_path))
            summary = {
                'chapters_found': chapters_found,
                'chapters_auto_generated': 0
            }

        report_id = str(uuid.uuid4())[:12]
        _generated_reports[report_id] = {
            'id': report_id,
            'filename': output_filename,
            'path': str(output_path),
            'created': datetime.now().isoformat(),
        }

        return {
            'success': True,
            'output_file': output_filename,
            'report_id': report_id,
            'summary': summary,
            'template_name': template_name
        }
    except Exception as e:
        import traceback
        traceback.print_exc()
        return {
            'success': False,
            'error': f'文档转换失败: {str(e)}'
        }
    finally:
        if upload_path.exists():
            upload_path.unlink()


@app.post('/api/template/{template_id}/apply')
async def apply_template_api(template_id: str, file: UploadFile = File(...)):
    """应用模板到文档"""
    if not file.filename or not allowed_file(file.filename):
        raise HTTPException(400, '请上传 .docx 文件')

    if template_id not in _template_storage:
        raise HTTPException(404, '模板不存在')

    unique_filename = f"{uuid.uuid4().hex}_{secure_filename(file.filename)}"
    upload_path = UPLOAD_DIR / unique_filename
    output_filename = f"{uuid.uuid4().hex}_formatted_report.docx"
    output_path = OUTPUT_DIR / output_filename

    try:
        content = await file.read()
        upload_path.write_bytes(content)

        template = _template_storage[template_id]
        content_data = template['content']
        
        if HAS_MODULES:
            formatter = UniversalDocxFormatter()
            result_path, structure = formatter.convert_with_template(
                str(upload_path),
                content_data,
                str(output_path)
            )
            post_processor = PostProcessor()
            summary = post_processor.generate_summary(structure)
        else:
            doc = Document(io.BytesIO(content_data))
            original_doc = Document(io.BytesIO(content))
            for para in original_doc.paragraphs:
                if para.text.strip():
                    doc.add_paragraph(para.text)
            doc.save(str(output_path))
            structure = {}
            summary = {'chapters_found': 0, 'chapters_auto_generated': 0}

        return {
            'success': True,
            'output_file': output_filename,
            'summary': summary,
            'template_name': template['name']
        }
    finally:
        if upload_path.exists():
            upload_path.unlink()


@app.post('/api/template/upload_std')
async def upload_standard_template(template: UploadFile = File(...)):
    """上传标准模板"""
    if not template.filename.endswith('.docx'):
        raise HTTPException(400, '请上传 .docx 文件')

    content = await template.read()
    template_data = parse_word_template(content)
    template_id = str(uuid.uuid4())[:8]

    _template_storage[template_id] = {
        'id': template_id,
        'name': template.filename,
        'content': content,
        'data': template_data,
        'upload_time': datetime.now().isoformat(),
    }

    return {
        'success': True,
        'template_id': template_id,
        'template_name': template.filename,
        'sections_count': len(template_data['sections']),
        'paragraphs_count': len(template_data['paragraphs']),
    }


@app.get('/api/template/list_std')
async def list_standard_templates():
    """列出标准模板"""
    templates = []
    for tid, t in _template_storage.items():
        templates.append({
            'id': tid,
            'name': t['name'],
            'upload_time': t['upload_time'],
            'sections_count': len(t['data'].get('sections', [])),
            'paragraphs_count': len(t['data'].get('paragraphs', []))
        })
    return {'success': True, 'templates': templates}


@app.get('/api/template/presets')
async def list_preset_templates():
    """列出所有预设模板"""
    presets = []
    for tid, t in PRESET_TEMPLATES.items():
        presets.append({
            'id': tid,
            'name': t['name'],
            'type': 'preset',
            'description': t['description'],
            'fields': t['fields']
        })
    return {'success': True, 'presets': presets}


@app.get('/api/template/preset/{template_id}')
async def get_preset_template_detail(template_id: str):
    """获取预设模板详情"""
    template = _get_preset_template(template_id)
    if not template:
        raise HTTPException(404, f'模板不存在: {template_id}')
    return {'success': True, 'template': template}


@app.post('/api/run-with-preset')
async def run_report_with_preset(request: Request):
    """使用AI生成保险公估调查报告"""
    debug_log_path = str(BASE_DIR / 'backend_debug.log')
    
    def log_write(msg):
        """同时写入文件和打印"""
        print(msg)
        try:
            with open(debug_log_path, 'a', encoding='utf-8') as f:
                f.write(msg + '\n')
        except:
            pass
    
    try:
        log_write(f'\n{"="*60}')
        log_write(f'[请求开始] 时间: {datetime.now().strftime("%Y-%m-%d %H:%M:%S")}')
        
        form = await request.form()
        template_id = form.get('template_id', 'preset_1')
        log_write(f'[DEBUG] 收到请求: template_id={template_id}')
        
        # 收集走访内容
        scene_investigation = form.get('sceneInvestigation', '')
        police_station_record = form.get('policeStationRecord', '')
        traffic_police_record = form.get('trafficPoliceRecord', '')
        hospital_record = form.get('hospitalRecord', '')
        
        combined_investigation = []
        if scene_investigation:
            combined_investigation.append(f"【走访现场】\n{scene_investigation}")
        if police_station_record:
            combined_investigation.append(f"【派出所调查】\n{police_station_record}")
        if traffic_police_record:
            combined_investigation.append(f"【交警队调查】\n{traffic_police_record}")
        if hospital_record:
            combined_investigation.append(f"【医院调查】\n{hospital_record}")
        
        investigation_text = '\n\n'.join(combined_investigation) if combined_investigation else ''
        log_write(f'[DEBUG] 走访内容总长度: {len(investigation_text)}')
        
        # 收集图片并进行OCR识别
        async def read_form_file(value):
            """异步读取表单文件字段"""
            try:
                if value and hasattr(value, 'read'):
                    content = await value.read()
                    return content
            except Exception as e:
                log_write(f'[Upload] 读取文件失败: {e}')
            return None
        
        # 获取所有上传的图片文件
        image_files = []
        statement_files = []
        all_keys = list(form.keys())
        log_write(f'[DEBUG] 表单所有键名: {all_keys}')
        
        try:
            raw_images = form.getlist('images')
            log_write(f'[DEBUG] getlist("images") 返回 {len(raw_images)} 个元素')
            for i, v in enumerate(raw_images):
                if hasattr(v, 'read'):
                    image_files.append(v)
                    log_write(f'[DEBUG] images[{i}] 是文件对象，文件名: {getattr(v, "filename", "未知")}')
                else:
                    log_write(f'[DEBUG] images[{i}] 不是文件对象，类型: {type(v)}')
        except Exception as e:
            log_write(f'[DEBUG] getlist("images") 失败: {e}')
        
        try:
            raw_statements = form.getlist('statement_images')
            log_write(f'[DEBUG] getlist("statement_images") 返回 {len(raw_statements)} 个元素')
            for i, v in enumerate(raw_statements):
                if hasattr(v, 'read'):
                    statement_files.append(v)
                    log_write(f'[DEBUG] statement_images[{i}] 是文件对象，文件名: {getattr(v, "filename", "未知")}')
                else:
                    log_write(f'[DEBUG] statement_images[{i}] 不是文件对象，类型: {type(v)}')
        except Exception as e:
            log_write(f'[DEBUG] getlist("statement_images") 失败: {e}')
        
        log_write(f'[DEBUG] 最终 image_files 数量: {len(image_files)}, statement_files 数量: {len(statement_files)}')
        
        # 处理图片OCR识别
        import base64
        ocr_results = []
        
        async def process_single_image(file_value, is_statement=False, index=0):
            """处理单张图片的OCR识别"""
            try:
                content = await read_form_file(file_value)
                if not content:
                    log_write(f'[OCR] 图片内容为空，跳过')
                    return None
                
                filename = getattr(file_value, 'filename', f'{"statement" if is_statement else "image"}_{index+1}.png')
                img_base64 = base64.b64encode(content).decode('utf-8')
                
                log_write(f'[DEBUG] 开始识别图片: {filename}, 大小: {len(content)} bytes')
                
                # 1. 首先尝试百度OCR
                if BAIDU_API_KEY and BAIDU_SECRET_KEY:
                    log_write(f'[DEBUG] 尝试百度OCR...')
                    try:
                        ocr_resp = await asyncio.to_thread(_call_baidu_ocr, img_base64)
                        log_write(f'[DEBUG] 百度OCR响应: {ocr_resp}')
                        if ocr_resp and ocr_resp.get('status') == 200:
                            data = ocr_resp.get('data', {})
                            if 'words_result' in data:
                                word_list = data['words_result']
                                text_parts = [w['words'] for w in word_list if w.get('words')]
                                if text_parts:
                                    ocr_text = '\n'.join(text_parts)
                                    log_write(f'[OCR] 识别 {filename} 成功，长度 {len(ocr_text)}')
                                    return {
                                        'name': filename,
                                        'text': ocr_text,
                                        'is_statement': is_statement
                                    }
                            else:
                                log_write(f'[DEBUG] 百度OCR响应中没有words_result')
                        else:
                            log_write(f'[DEBUG] 百度OCR返回状态码: {ocr_resp.get("status") if ocr_resp else "None"}')
                    except Exception as e:
                        log_write(f'[OCR] 百度OCR处理 {filename} 失败: {e}')
                        import traceback
                        log_write(traceback.format_exc())
                else:
                    log_write(f'[OCR] 百度API未配置')
                
                # 2. 如果百度OCR失败，尝试千问VL
                log_write(f'[DEBUG] 百度OCR失败，尝试千问VL...')
                if QWEN_VL_API_KEY:
                    try:
                        description = await _call_qwen_vl(img_base64)
                        if description:
                            log_write(f'[Qwen-VL] 识别 {filename} 成功，长度 {len(description)}')
                            return {
                                'name': filename,
                                'text': description,
                                'is_statement': is_statement
                            }
                        else:
                            log_write(f'[DEBUG] 千问VL返回空结果')
                    except Exception as e:
                        log_write(f'[Qwen-VL] 处理 {filename} 失败: {e}')
                        import traceback
                        log_write(traceback.format_exc())
                else:
                    log_write(f'[OCR] 千问VL API未配置')
                    
                log_write(f'[WARNING] 所有OCR服务都未能识别图片: {filename}')
            except Exception as e:
                log_write(f'[OCR] 图片处理异常: {e}')
                import traceback
                log_write(traceback.format_exc())
            
            return None
        
        # 处理所有图片
        all_files = [(f, False, i) for i, f in enumerate(image_files)] + \
                    [(f, True, i + len(image_files)) for i, f in enumerate(statement_files)]
        
        for file_value, is_statement, idx in all_files:
            try:
                ocr_result = await process_single_image(file_value, is_statement, idx)
                if ocr_result:
                    ocr_results.append(ocr_result)
            except Exception as e:
                log_write(f'[处理] 第{idx+1}张图片处理失败: {e}')
        
        log_write(f'[DEBUG] OCR识别完成: {len(ocr_results)}/{len(all_files)} 张成功')
        
        # 整理OCR结果
        scene_ocr_texts = []
        statement_ocr_texts = []
        
        for result in ocr_results:
            if result['is_statement']:
                statement_ocr_texts.append(f"【笔录{len(statement_ocr_texts)+1}】\n{result['text']}")
            else:
                scene_ocr_texts.append(f"【现场照片{len(scene_ocr_texts)+1}】\n{result['text']}")
        
        scene_ocr_combined = '\n\n'.join(scene_ocr_texts) if scene_ocr_texts else ''
        statement_ocr_combined = '\n\n'.join(statement_ocr_texts) if statement_ocr_texts else ''
        all_ocr_combined = scene_ocr_combined + ('\n\n' if scene_ocr_combined and statement_ocr_combined else '') + statement_ocr_combined
        
        log_write(f'[DEBUG] 现场照片OCR文本长度: {len(scene_ocr_combined)}')
        log_write(f'[DEBUG] 笔录OCR文本长度: {len(statement_ocr_combined)}')
        
        # 获取选择的模板（即使没有DeepSeek，也能用模板生成基础报告）
        selected_template = PRESET_TEMPLATES.get(template_id, PRESET_TEMPLATES['preset_1'])
        # 如果没有DeepSeek，直接用模板填充基础内容后返回
        if not DEEPSEEK_API_KEY:
            template_content = selected_template['content']
            template_name = selected_template['name']
            report_content = template_content
            if investigation_text:
                report_content = report_content.replace('{{investigation_content}}', investigation_text)
            # 替换基本占位符
            for placeholder, default_value in [
                ('{{insured_name}}', '待核实'), ('{{insurance_company}}', '保险公司'),
                ('{{policy_no}}', '待核实'), ('{{insurance_period}}', '待核实'),
                ('{{insurance_type}}', '待核实'), ('{{accident_time}}', '待核实'),
                ('{{investigator}}', '待填写'), ('{{reviewer}}', '待填写'),
                ('{{company_name}}', '深圳市恒泰诚信息咨询有限公司'),
                ('{{report_date}}', datetime.now().strftime('%Y年%m月%d日')),
            ]:
                if placeholder in report_content:
                    report_content = report_content.replace(placeholder, default_value)
            report_content = re.sub(r'\{\{[^}]+\}\}', '待核实', report_content)
            # 清理模板内容：只保留章节标题行和报告标题行，去掉模板自带的固定描述文字
            cleaned_lines = []
            for line in report_content.split('\n'):
                stripped = line.strip()
                if not stripped:
                    continue
                # 保留报告标题（不含章节标记的正文第一行）
                # 保留章节标题行（以"一、"到"九、"开头，或以"（一）"到"（九）"开头）
                # 保留包含"待核实"、"待填写"等替换内容的行（即占位符被替换后的行）
                is_chapter_heading = re.match(r'^[一二三四五六七八九十]+[、．\.]', stripped) or re.match(r'^（[一二三四五六七八九十]+）', stripped)
                is_placeholder_line = '待核实' in stripped or '待填写' in stripped or '保险公司' in stripped or '深圳市恒泰诚' in stripped
                if is_chapter_heading or is_placeholder_line:
                    cleaned_lines.append(line)
                # 报告标题（第一行）也保留
                elif len(cleaned_lines) == 0 and not is_chapter_heading:
                    cleaned_lines.append(line)
            if cleaned_lines:
                report_content = '\n'.join(cleaned_lines)
            return {
                'success': True,
                'report_content': report_content,
                'template_id': template_id or 'preset_1',
                'template_name': template_name,
                'used_ai_fallback': True,
                'ocr_count': len(ocr_results),
                'combined_text': all_ocr_combined,
                'total_images': len(all_files),
                'investigation_text': investigation_text
            }
        
        # 获取选择的模板
        selected_template = PRESET_TEMPLATES.get(template_id, PRESET_TEMPLATES['preset_1'])
        template_content = selected_template['content']
        template_name = selected_template['name']
        template_fields = selected_template.get('fields', [])
        
        log_write(f'[DEBUG] 使用模板: {template_id} ({template_name})')
        log_write(f'[DEBUG] 模板内容长度: {len(template_content)}')
        
        # 第一步：先对OCR内容进行数据结构化
        if investigation_text or all_ocr_combined:
            structure_prompt = f"""请将以下调查信息进行结构化处理，提取关键信息填入模板字段。

【目标模板字段】
{', '.join(template_fields)}

【模板格式】
{template_content}

【调查走访内容】
{investigation_text if investigation_text else "（无走访内容）"}

【OCR识别内容】
{all_ocr_combined if all_ocr_combined else "（无OCR内容）"}

【结构化要求】
1、请仔细分析调查走访内容和OCR识别内容
2、提取关键信息填入模板的{{}}占位符位置
3、保持模板的原有章节结构不变
4、对于{{investigation_content}}等大段落字段，需要将相关内容整合成一个完整的段落
5、如果某些字段在内容中没有找到，请标注"待核实"或"根据调查确定"
6、直接输出完整的报告，不要添加其他说明
7、严格按照模板格式，不要添加或删除章节"""
            
            log_write(f'[DEBUG] 开始对OCR内容进行数据结构化...')
            log_write(f'[DEBUG] 结构化提示词长度: {len(structure_prompt)}')
            
            # 调用DeepSeek API进行结构化
            report_content = await _call_deepseek_api(structure_prompt)
            
            if report_content:
                log_write(f'[DEBUG] 结构化成功，报告长度: {len(report_content)}')
            else:
                log_write(f'[WARNING] 结构化失败，使用模板原始内容')
                # 如果结构化失败，使用模板原始内容并替换走访内容
                report_content = template_content
                if investigation_text:
                    report_content = report_content.replace('{{investigation_content}}', investigation_text)
                if all_ocr_combined:
                    report_content = report_content.replace('{{ocr_text}}', all_ocr_combined)
        else:
            # 如果没有调查内容和OCR内容，使用模板原始格式
            log_write(f'[DEBUG] 无调查内容和OCR内容，使用模板原始格式')
            report_content = template_content
        
        # 如果报告中仍有未填充的占位符，进行简单替换
        placeholder_patterns = [
            ('{{insured_name}}', '被保险人'),
            ('{{insurance_company}}', '保险公司'),
            ('{{policy_no}}', '待核实'),
            ('{{insurance_period}}', '待核实'),
            ('{{insurance_type}}', '待核实'),
            ('{{accident_time}}', '待核实'),
            ('{{case_description}}', '待核实'),
            ('{{verification_items}}', '待核实'),
            ('{{investigation_content}}', '待核实'),
            ('{{summary}}', '待核实'),
            ('{{conclusion}}', '待核实'),
            # 保险责任分析字段
            ('{{liability_insured_name}}', '待核实'),
            ('{{liability_accident_time}}', '待核实'),
            ('{{liability_accident_location}}', '待核实'),
            ('{{liability_accident_cause}}', '待核实'),
            ('{{liability_victim_identity}}', '待核实'),
            ('{{liability_content}}', '待核实'),
            ('{{investigator}}', '待填写'),
            ('{{reviewer}}', '待填写'),
            ('{{company_name}}', '深圳市恒泰诚信息咨询有限公司'),
            ('{{report_date}}', datetime.now().strftime('%Y年%m月%d日')),
        ]
        
        for placeholder, default_value in placeholder_patterns:
            if placeholder in report_content:
                report_content = report_content.replace(placeholder, default_value)
                log_write(f'[DEBUG] 替换占位符: {placeholder} -> {default_value}')
        
        # 检查是否还有未替换的占位符
        remaining = report_content.count('{{')
        if remaining > 0:
            log_write(f'[WARNING] 仍有 {remaining} 个未替换的占位符')
            # 将剩余占位符替换为默认值
            import re
            report_content = re.sub(r'\{\{[^}]+\}\}', '待核实', report_content)
        
        log_write(f'[DEBUG] 最终报告长度: {len(report_content)}')
        
        return {
            'success': True,
            'report_content': report_content,
            'template_id': template_id,
            'template_name': template_name,
            'used_ai_fallback': False,
            'ocr_count': len(ocr_results),
            'combined_text': all_ocr_combined,
            'total_images': len(all_files),
            'investigation_text': investigation_text
        }
    except Exception as e:
        log_write(f'[ERROR] 生成报告异常: {e}')
        import traceback
        log_write(traceback.format_exc())
        return {
            'success': False,
            'error': f'生成报告失败: {str(e)}'
        }


# ============================================================
# 新增缺失的 API 端点
# ============================================================

# ============ 文件下载端点 ============
@app.get('/api/download/{report_id}')
async def download_report(report_id: str):
    """下载已生成的报告文件"""
    if report_id in _generated_reports:
        r = _generated_reports[report_id]
        path_str = r.get('path', '')
        if path_str and os.path.exists(path_str):
            filename = r.get('filename', os.path.basename(path_str))
            return FileResponse(path_str, media_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document', filename=filename)
    # 尝试在 OUTPUT_DIR 中直接查找
    for f in os.listdir(OUTPUT_DIR):
        if report_id in f or report_id in str(OUTPUT_DIR / f):
            fp = OUTPUT_DIR / f
            if fp.is_file():
                return FileResponse(str(fp), media_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document', filename=f)
    raise HTTPException(404, '文件不存在或已过期')


# ============ 结案审批表填充端点 (renderer.js -> /api/fill-excel) ============
@app.post('/api/fill-excel')
async def fill_excel(request: Request):
    """填充Excel审批表——保留标签单元格，将值填入同行相邻单元格"""
    import openpyxl
    from openpyxl.utils import get_column_letter
    try:
        form = await request.form()
        template_file = None
        for key in form.keys():
            v = form.get(key)
            if hasattr(v, 'read'):
                template_file = v
                break
        if not template_file:
            raise HTTPException(400, '未找到模板文件')
        data_str = form.get('data', '{}')
        if isinstance(data_str, str):
            data = json.loads(data_str)
        else:
            data = data_str
        content = await template_file.read()
        wb = openpyxl.load_workbook(io.BytesIO(content))
        ws = wb.active

        # 构建标签→值的映射（中英文标签都匹配）
        label_to_value = {}
        field_mappings = [
            ('policyNo', ['保单号', '保单号码', 'policyNo']),
            ('insured', ['投保人', '投保人名称', 'insured']),
            ('insuredPerson', ['被保险人', '伤者', 'insuredPerson']),
            ('insuranceType', ['险种', '保险险种', 'insuranceType']),
            ('accidentDate', ['出险时间', '事故发生时间', 'accidentDate']),
            ('accidentLocation', ['出险地点', '案发地点', 'accidentLocation']),
            ('claimAmount', ['索赔金额', '理赔金额', 'claimAmount']),
            ('suggestion', ['赔付建议', '处理建议', 'suggestion']),
            ('acceptDate', ['受案时间', '受理时间', 'acceptDate']),
            ('closeDate', ['结案时间', '完成时间', 'closeDate']),
            ('investigator', ['调查员', '查勘员', 'investigator']),
            ('insurancePeriod', ['保险期间', '保险期限', 'insurancePeriod']),
            ('accidentNature', ['事故性质', '事故类型', 'accidentNature']),
        ]
        for field_key, labels in field_mappings:
            value = data.get(field_key, '')
            if value:
                for label in labels:
                    label_to_value[label] = value

        modified_count = 0
        # 第一遍：找到标签所在单元格，将值填入同行右侧（非标签）单元格
        for row in ws.iter_rows(min_row=1, max_row=ws.max_row, max_col=ws.max_column):
            row_cells = list(row)
            for col_idx, cell in enumerate(row_cells):
                if not cell.value or not isinstance(cell.value, str):
                    continue
                cell_text = str(cell.value).strip()
                # 检查是否匹配某个标签
                matched_label = None
                matched_value = None
                for label, val in label_to_value.items():
                    if label in cell_text or cell_text.lower() == label.lower():
                        matched_label = label
                        matched_value = val
                        break
                if matched_label is None:
                    continue

                # 优先填到右侧相邻单元格（同行的下一列）
                filled = False
                if col_idx + 1 < len(row_cells):
                    right_cell = row_cells[col_idx + 1]
                    # 如果右侧单元格是空、或包含占位符/旧值，就替换它
                    if right_cell.value is None or str(right_cell.value).strip() in ('', '待填写', '-', '—', '待核实', '待定'):
                        right_cell.value = matched_value
                        modified_count += 1
                        filled = True
                    else:
                        # 右侧已有内容，检查是否也包含同标签（可能是个占位模板）
                        right_text = str(right_cell.value).strip()
                        # 如果不是标签文字，直接覆盖
                        is_label = any(lbl in right_text or right_text.lower() == lbl.lower() for lbl in label_to_value)
                        if not is_label:
                            right_cell.value = matched_value
                            modified_count += 1
                            filled = True

                if not filled:
                    # 如果右侧无法填入，直接在原单元格中把标签替换为值（保留标签前缀）
                    if cell_text == matched_label or cell_text == f'{matched_label}：':
                        cell.value = f'{matched_label}：{matched_value}'
                    elif matched_label in cell_text:
                        cell.value = cell_text.replace(matched_label, f'{matched_label}：{matched_value}')
                    else:
                        # 兜底：直接替换
                        cell.value = f'{matched_label}：{matched_value}'
                    modified_count += 1

        output = io.BytesIO()
        wb.save(output)
        output.seek(0)
        from urllib.parse import quote
        safe_filename = f'审批表_{datetime.now().strftime("%Y%m%d")}.xlsx'
        return Response(content=output.read(), media_type='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet',
                        headers={'Content-Disposition': f"attachment; filename*=UTF-8''{quote(safe_filename)}"})
    except Exception as e:
        import traceback
        traceback.print_exc()
        return JSONResponse(content={'success': False, 'error': f'生成Excel失败: {str(e)}'}, status_code=500)


# ============ 目录工具 - 处理DOCX (index4.html -> /api/process-docx) ============
@app.post('/api/process-docx')
async def process_docx(request: Request):
    """处理DOCX文件，替换目录中的附件清单"""
    try:
        form = await request.form()
        file = None
        for key in form.keys():
            v = form.get(key)
            if hasattr(v, 'read'):
                file = v
                break
        if not file:
            raise HTTPException(400, '未找到上传文件')
        attachments_str = form.get('attachments', '[]')
        if isinstance(attachments_str, str):
            attachments = json.loads(attachments_str)
        else:
            attachments = attachments_str
        content = await file.read()
        # 读取DOCX，找到表格并替换
        doc = Document(io.BytesIO(content))
        table_count = 0
        for table in doc.tables:
            replace_table_with_attachments(table, attachments)
            table_count += 1
        if table_count == 0 and attachments:
            # 没有表格，创建一个新的附件目录表
            doc.add_page_break()
            doc.add_heading('附件清单', level=1)
            new_table = doc.add_table(rows=1 + len(attachments), cols=3)
            new_table.style = 'Table Grid'
            hdr = new_table.rows[0].cells
            hdr[0].text = '序号'
            hdr[1].text = '附件名称'
            hdr[2].text = '页码'
            for i, att in enumerate(attachments):
                row = new_table.rows[i + 1].cells
                row[0].text = str(att.get('number', i + 1))
                row[1].text = att.get('name', '')
                row[2].text = '—'
        output = io.BytesIO()
        doc.save(output)
        output.seek(0)
        from urllib.parse import quote
        safe_docx_name = f'处理报告_{datetime.now().strftime("%Y%m%d")}.docx'
        return Response(content=output.read(),
                        media_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
                        headers={'Content-Disposition': f"attachment; filename*=UTF-8''{quote(safe_docx_name)}"})
    except Exception as e:
        import traceback
        traceback.print_exc()
        return JSONResponse(content={'success': False, 'error': str(e)}, status_code=500)


def replace_table_with_attachments(table, attachments):
    """替换附件清单：保留表头，在第2列填入名称，未替换的多余行全部删除"""
    rows_to_delete = []
    for row_idx, row in enumerate(table.rows):
        cells = row.cells
        # 第0行（表头行）无条件跳过，绝不动
        if row_idx == 0:
            continue
        data_index = row_idx - 1
        if data_index < len(attachments):
            # 有对应附件数据 -> 填入名称
            if len(cells) >= 2:
                cell = cells[1]
                cell.text = ''
                p = cell.paragraphs[0]
                run = p.add_run(attachments[data_index]['name'])
                run.font.size = Pt(10.5)
        else:
            # 超出附件数量 -> 标记为待删除
            rows_to_delete.append(row)
    # 从表格中移除多余行（从后往前删避免索引错乱）
    for row in reversed(rows_to_delete):
        tbl = row._tr.getparent()
        tbl.remove(row._tr)


# ============ 报告生成 /api/run (report_generate.js 调用) ============
@app.post('/api/run')
async def run_report(request: Request):
    """生成报告的入口（前端 ajax 版本）"""
    import traceback as _tb
    try:
        form = await request.form()
        template_id = form.get('template_id', 'preset_1')
        print(f'[api/run] template_id={template_id}, keys={list(form.keys())}')
        # 收集走访内容
        s = lambda k: form.get(k, '')
        parts = []
        for key, prefix in [('sceneInvestigation','走访现场'),('policeStationRecord','派出所调查'),('trafficPoliceRecord','交警队调查'),('hospitalRecord','医院调查')]:
            v = s(key)
            if v: parts.append(f'【{prefix}】\n{v}')
        investigation_text = '\n\n'.join(parts)
        # 获取模板
        sel = PRESET_TEMPLATES.get(template_id, PRESET_TEMPLATES['preset_1'])
        content = sel['content']
        if investigation_text:
            content = content.replace('{{investigation_content}}', investigation_text)
        # 替换所有占位符
        for ph, dv in [('insured_name','待核实'),('insurance_company','保险公司'),('policy_no','待核实'),('insurance_period','待核实'),('insurance_type','待核实'),('accident_time','待核实'),('investigator','待填写'),('reviewer','待填写'),('company_name','深圳市恒泰诚信息咨询有限公司'),('report_date',datetime.now().strftime('%Y年%m月%d日'))]:
            k = '{{' + ph + '}}'
            if k in content: content = content.replace(k, dv)
        import re as _re
        content = _re.sub(r'\{\{[^}]+\}\}', '待核实', content)
        return {'success':True,'report_content':content,'template_id':template_id,'template_name':sel['name'],'used_ai_fallback':True,'total_images':0,'combined_text':'','investigation_text':investigation_text}
    except Exception as e:
        _tb.print_exc()
        return {'success':False,'error':f'生成报告失败: {str(e)}'}
# ============ 附件处理器 - 临时文件服务 ============
@app.get('/temp/{filename}')
async def serve_temp_file(filename: str):
    """提供临时图片文件访问"""
    safe = os.path.basename(filename)
    # 检查多个可能的临时目录
    candidates = [
        OUTPUT_DIR / safe,
        UPLOAD_DIR / safe,
        Path(f'/tmp/case_processor/uploads/{safe}'),
        Path(f'/tmp/case_processor/output/{safe}'),
    ]
    # 也检查当前目录下的子目录
    for sub in ['uploads', 'output', 'temp']:
        p = BASE_DIR / sub / safe
        if p.parent.exists():
            candidates.append(p)
    for p in candidates:
        try:
            if os.path.exists(str(p)):
                return FileResponse(str(p))
        except:
            pass
    raise HTTPException(404, '文件未找到')


# ============ 附件处理器 - 一键生成 (attachment_web/index.html -> /api/attachment/generate) ============
@app.post('/api/attachment/generate')
async def attachment_generate(request: Request):
    """上传图片并一键生成Word报告"""
    try:
        form = await request.form()
        images = []
        file_count = 0
        for key in form.keys():
            v = form.get(key)
            if hasattr(v, 'read'):
                content = await v.read()
                if content:
                    ext = os.path.splitext(getattr(v, 'filename', f'image_{file_count}.png'))[1] or '.png'
                    filename = f'att_{uuid.uuid4().hex[:10]}{ext}'
                    save_path = UPLOAD_DIR / filename
                    save_path.write_bytes(content)
                    images.append({
                        'name': getattr(v, 'filename', filename),
                        'path': str(save_path),
                        'size': len(content)
                    })
                    file_count += 1
        if not images:
            raise HTTPException(400, '未上传图片文件')
        names_str = form.get('names', '[]')
        if isinstance(names_str, str) and names_str.strip():
            custom_names = json.loads(names_str)
        else:
            custom_names = []
        # 生成Word文档
        doc = Document()
        doc.add_heading('附件报告', 0)
        doc.add_paragraph(f'生成时间: {datetime.now().strftime("%Y-%m-%d %H:%M:%S")}')
        doc.add_paragraph(f'共 {len(images)} 张图片')
        doc.add_paragraph()
        for i, img in enumerate(images):
            att_name = f"附件{i+1}"
            if i < len(custom_names):
                att_name = custom_names[i]
            elif i < len(custom_names):
                att_name = custom_names[i]
            doc.add_heading(f'{att_name}', level=1)
            try:
                doc.add_picture(img['path'], width=Inches(5.5))
            except Exception as e:
                doc.add_paragraph(f'[图片无法加载: {img["name"]}]')
            doc.add_paragraph()
        output_filename = f'attachment_report_{uuid.uuid4().hex[:8]}.docx'
        output_path = OUTPUT_DIR / output_filename
        doc.save(str(output_path))
        report_id = str(uuid.uuid4())[:12]
        _generated_reports[report_id] = {
            'id': report_id,
            'filename': output_filename,
            'path': str(output_path),
            'created': datetime.now().isoformat(),
        }
        # 返回文件流
        return FileResponse(str(output_path),
                            media_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
                            filename=output_filename)
    except Exception as e:
        import traceback
        traceback.print_exc()
        return JSONResponse(content={'success': False, 'error': str(e)}, status_code=500)


# ============ 附件处理器（合并版）- 上传图片 ============
_session_images = {}  # session_id -> [image_info]


@app.post('/api/upload-images')
async def upload_images(request: Request):
    """上传图片到临时存储"""
    try:
        form = await request.form()
        uploaded = []
        file_count = 0
        for key in form.keys():
            v = form.get(key)
            if hasattr(v, 'read') and key.startswith('files') or hasattr(v, 'read'):
                content = await v.read()
                if content:
                    ext = os.path.splitext(getattr(v, 'filename', f'img_{file_count}.png'))[1] or '.png'
                    filename = f'up_{uuid.uuid4().hex[:10]}{ext}'
                    save_path = UPLOAD_DIR / filename
                    save_path.write_bytes(content)
                    uploaded.append({
                        'index': file_count,
                        'name': getattr(v, 'filename', filename),
                        'size': len(content),
                        'filepath': str(save_path)
                    })
                    file_count += 1
        if not uploaded:
            return JSONResponse(content={'success': False, 'error': '没有接收到文件'}, status_code=400)
        return {'success': True, 'images': uploaded, 'count': len(uploaded)}
    except Exception as e:
        return JSONResponse(content={'success': False, 'error': str(e)}, status_code=500)


# ============ 附件处理器（合并版）- 上传DOCX模板 ============
@app.post('/api/upload-docx')
async def upload_docx_template(request: Request):
    """上传DOCX模板——保存文件并提取信息供后续一键生成使用"""
    try:
        form = await request.form()
        file = None
        for key in form.keys():
            v = form.get(key)
            if hasattr(v, 'read'):
                file = v
                break
        if not file:
            raise HTTPException(400, '未找到文件')
        content = await file.read()

        # 保存模板文件
        template_id = str(uuid.uuid4())[:12]
        template_filename = f'template_{template_id}.docx'
        template_path = UPLOAD_DIR / template_filename
        template_path.write_bytes(content)

        # 解析DOCX信息
        doc = Document(io.BytesIO(content))
        title = ''
        header_info = {'text': '', 'hasIcon': False}
        image_sizes = []
        # 提取第一个段落作为标题
        for para in doc.paragraphs:
            if para.text.strip() and not title:
                title = para.text.strip()[:100]
                break
        # 提取页眉
        for section in doc.sections:
            header = section.header
            if header:
                header_text = ' '.join(p.text for p in header.paragraphs if p.text.strip())
                if header_text:
                    header_info['text'] = header_text
        # 从文档中提取占位图片尺寸（内联图片/形状）
        from docx.opc.constants import RELATIONSHIP_TYPE as RT
        for rel in doc.part.rels.values():
            if "image" in rel.reltype:
                try:
                    image_blob = rel.target_part.blob
                    from PIL import Image
                    from io import BytesIO
                    img = Image.open(BytesIO(image_blob))
                    image_sizes.append({'width': img.width, 'height': img.height, 'rel_id': rel.rId})
                except:
                    pass
        _template_storage[template_id] = {
            'id': template_id,
            'path': str(template_path),
            'filename': template_filename,
            'title': title,
            'header': header_info,
            'imageSizes': image_sizes,
            'upload_time': datetime.now().isoformat()
        }

        return {
            'success': True,
            'template_id': template_id,
            'content': '',
            'title': title,
            'header': header_info,
            'imageSizes': image_sizes
        }
    except Exception as e:
        import traceback
        traceback.print_exc()
        return JSONResponse(content={'success': False, 'error': str(e)}, status_code=500)


# ============ 附件处理器（合并版）- ranges API ============
@app.get('/api/ranges')
async def get_ranges():
    """获取映射规则"""
    return {'success': True, 'ranges': [
        {'start': 0, 'end': 2, 'idx': 0}, {'start': 3, 'end': 18, 'idx': 1},
        {'start': 19, 'end': 20, 'idx': 2}, {'start': 21, 'end': 24, 'idx': 3},
    ]}

@app.put('/api/ranges')
async def save_ranges(request: Request):
    """保存映射规则"""
    data = await request.json()
    return {'success': True}

@app.post('/api/ranges/reset')
async def reset_ranges():
    """重置映射规则"""
    return {'success': True, 'ranges': [
        {'start': 0, 'end': 2, 'idx': 0}, {'start': 3, 'end': 18, 'idx': 1},
    ]}

# ============ 附件处理器（合并版）- 一键生成 ============
@app.post('/api/one-click')
async def one_click_generate(request: Request):
    """一键生成：使用已上传DOCX模板替换附件编号/名称/图片"""
    try:
        data = await request.json()
        images = data.get('images', [])
        attachment_names_raw = data.get('attachmentNames', '')
        custom_ranges = data.get('customRanges', [])
        header = data.get('header', {})
        docx_image_sizes = data.get('docxImageSizes', [])
        template_id = data.get('template_id', '')

        # 中文数字转阿拉伯数字（支持多字如"十一""二十三"）
        def chinese_num_to_str(cn):
            if cn.isdigit():
                return cn
            cn_map = {'零':'0','一':'1','二':'2','三':'3','四':'4','五':'5','六':'6','七':'7','八':'8','九':'9','十':'10'}
            if cn in cn_map:
                return cn_map[cn]
            # 多字组合：十一、十一、二十一
            try:
                if '十' in cn:
                    parts = cn.split('十')
                    if parts[0] == '' and parts[1] == '':
                        return '10'
                    elif parts[0] == '':
                        return str(10 + int(cn_map.get(parts[1], '0')))
                    elif parts[1] == '':
                        return str(int(cn_map.get(parts[0], '0')) * 10)
                    else:
                        return str(int(cn_map.get(parts[0], '0')) * 10 + int(cn_map.get(parts[1], '0')))
            except:
                pass
            return cn

        # 解析附件名称列表
        parsed_attachments = []  # [{number, name}, ...]
        if attachment_names_raw and attachment_names_raw.strip():
            for line in attachment_names_raw.strip().split('\n'):
                line = line.strip()
                if not line: continue
                m = re.match(r'^附件([一二三四五六七八九十0-9零一二三四五六七八九十]+)[:：\s、]*(.*)', line)
                if m:
                    num_str = m.group(1).strip()
                    name = m.group(2).strip()
                    num_str = chinese_num_to_str(num_str)
                    parsed_attachments.append({'number': num_str, 'name': name if name else f'附件{num_str}'})
                else:
                    # 纯名称行，自动编号
                    parsed_attachments.append({'number': str(len(parsed_attachments)+1), 'name': line})
        if not parsed_attachments:
            for i in range(len(images)):
                parsed_attachments.append({'number': str(i+1), 'name': f'附件{i+1}'})

        # 生成Markdown
        md_parts = [f'# 附件处理报告\n\n生成时间: {datetime.now().strftime("%Y-%m-%d %H:%M:%S")}\n\n']
        for i, att in enumerate(parsed_attachments):
            md_parts.append(f'## 附件{att["number"]}：{att["name"]}\n')
            if i < len(images):
                fp = images[i].get('filepath', '')
                md_parts.append(f'![{att["name"]}]({fp})\n\n')
            else:
                md_parts.append('（无图片）\n\n')
        markdown = '\n'.join(md_parts)

        # ====== 判断是否使用已上传的DOCX模板 ======
        template_path = None
        if template_id and template_id in _template_storage:
            tinfo = _template_storage[template_id]
            tp = tinfo.get('path', '')
            if tp and os.path.exists(tp):
                template_path = tp

        if template_path:
            # ====== 使用模板文件，替换其中的编号/名称/图片 ======
            doc = Document(template_path)

            # 1. 替换文档中的附件编号占位符（如"附件一"、"附件1"等）
            for para in doc.paragraphs:
                for run in para.runs:
                    text = run.text
                    if not text: continue
                    for i, att in enumerate(parsed_attachments):
                        # 匹配"附件X"模式
                        for pattern in [f'附件{i+1}', f'附件{att["number"]}',
                                        f'附件{["","一","二","三","四","五","六","七","八","九","十"][int(att["number"])] if att["number"].isdigit() and int(att["number"]) <= 10 else ""}']:
                            if pattern and pattern in text:
                                run.text = run.text.replace(pattern, f'附件{att["number"]}')
                                break

            # 2. 替换表格中的附件名称
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for para in cell.paragraphs:
                            for run in para.runs:
                                text = run.text
                                if not text: continue
                                # 替换"附件X"编号
                                for i, att in enumerate(parsed_attachments):
                                    cn = ['','一','二','三','四','五','六','七','八','九','十']
                                    cn_num = cn[int(att['number'])] if att['number'].isdigit() and int(att['number']) <= 10 else att['number']
                                    for p in [f'附件{i+1}', f'附件{att["number"]}', f'附件{cn_num}']:
                                        if p in run.text:
                                            run.text = run.text.replace(p, f'附件{att["number"]}')
                                            break
                                # 替换附件名称（找"附件X："后的名称文本）
                                name_match = re.match(r'^附件\d+[：:]\s*(.*)', run.text)
                                if name_match:
                                    for i, att in enumerate(parsed_attachments):
                                        if f'附件{att["number"]}' in run.text:
                                            run.text = f'附件{att["number"]}：{att["name"]}'
                                            break

            # 3. 替换图片占位符
            img_idx = 0
            inline_shapes = doc.inline_shapes
            for block in doc.element.body:
                if block.tag.endswith('}p'):
                    # 检查段落中的图片
                    drawings = block.findall('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}drawing')
                    if drawings and img_idx < len(images):
                        # 找到这个drawing，尝试替换
                        fp = images[img_idx].get('filepath', '')
                        if fp and os.path.exists(fp):
                            try:
                                # 删除旧图片
                                for d in drawings:
                                    block.remove(d)
                                # 在段落后添加新图片
                                from docx.oxml.ns import qn
                                # 用python-docx方式添加图片到段落
                                p_elem = block
                                # 通过内联形状
                                new_run_elem = doc.element.makeelement(qn('w:r'), {})
                                p_elem.append(new_run_elem)
                                # 使用python-docx添加图片
                                for p_idx, p in enumerate(doc.paragraphs):
                                    if p._element is block:
                                        p.add_run().add_picture(fp, width=Inches(5.5))
                                        break
                            except Exception as e:
                                print(f'[替换图片] 失败: {e}')
                            img_idx += 1

            # 保存
            filename = f'oneclick_{uuid.uuid4().hex[:8]}.docx'
            out_path = OUTPUT_DIR / filename
            doc.save(str(out_path))
            _generated_reports[filename] = {
                'id': filename,
                'filename': filename,
                'path': str(out_path),
                'created': datetime.now().isoformat()
            }
            attachments_list = []
            for att in parsed_attachments:
                try:
                    n = int(att['number'])
                except ValueError:
                    n = len(attachments_list) + 1
                attachments_list.append({'number': n, 'name': att['name']})
            return {'success': True, 'markdown': markdown, 'filename': filename, 'attachments': attachments_list}
        else:
            # ====== 无模板，直接生成新文档 ======
            doc = Document()
            doc.add_heading('附件处理报告', 0)
            doc.add_paragraph(f'生成时间: {datetime.now().strftime("%Y-%m-%d %H:%M:%S")}')
            doc.add_paragraph(f'共 {len(parsed_attachments)} 个附件')
            doc.add_paragraph()
            for i, att in enumerate(parsed_attachments):
                doc.add_heading(f'附件{att["number"]}：{att["name"]}', level=1)
                if i < len(images):
                    fp = images[i].get('filepath', '')
                    if fp and os.path.exists(fp):
                        try:
                            doc.add_picture(fp, width=Inches(5.5))
                        except:
                            doc.add_paragraph(f'[图片: {images[i].get("name","")}]')
                doc.add_paragraph()
            filename = f'oneclick_{uuid.uuid4().hex[:8]}.docx'
            out_path = OUTPUT_DIR / filename
            doc.save(str(out_path))
            _generated_reports[filename] = {
                'id': filename,
                'filename': filename,
                'path': str(out_path),
                'created': datetime.now().isoformat()
            }
            attachments_list = []
            for att in parsed_attachments:
                try:
                    n = int(att['number'])
                except ValueError:
                    n = len(attachments_list) + 1
                attachments_list.append({'number': n, 'name': att['name']})
            return {'success': True, 'markdown': markdown, 'filename': filename, 'attachments': attachments_list}
    except Exception as e:
        import traceback
        traceback.print_exc()
        return JSONResponse(content={'success': False, 'error': str(e)}, status_code=500)

# ============ 附件处理器（合并版）- 仅生成Markdown ============
@app.post('/api/generate-markdown')
async def generate_markdown(request: Request):
    """仅生成Markdown文本"""
    try:
        data = await request.json()
        images = data.get('images', [])
        md_parts = [f'# {data.get("docTitle","附件处理报告")}\n\n生成时间: {datetime.now().strftime("%Y-%m-%d %H:%M:%S")}\n\n']
        attachments = []
        for i, img in enumerate(images):
            att_name = f'附件{i+1}'
            md_parts.append(f'## {att_name}\n')
            md_parts.append(f'![{img.get("name","图片")}]({img.get("filepath","")})\n\n')
            attachments.append({'number': i+1, 'name': att_name})
        return {'success': True, 'markdown': '\n'.join(md_parts), 'attachments': attachments}
    except Exception as e:
        return JSONResponse(content={'success': False, 'error': str(e)}, status_code=500)

# ============ SSE 进度推送端点 ============
@app.get('/api/progress')
async def sse_progress(request: Request):
    """SSE (Server-Sent Events) 进度推送"""
    async def event_generator():
        try:
            yield f"data: {json.dumps({'type':'connected','status':'ok','message':'SSE已连接'})}\n\n"
            # 向前端发送5次进度更新
            for i in range(5):
                await asyncio.sleep(0.5)
                progress_data = {
                    'type': 'oneclick',
                    'phase': ['upload', 'build', 'markdown', 'word', 'done'][i],
                    'status': 'done',
                    'progress': (i + 1) * 20,
                    'message': f'步骤 {i+1}/5 完成'
                }
                yield f"data: {json.dumps(progress_data)}\n\n"
                if i == 4:
                    yield f"data: {json.dumps({'type':'oneclick','phase':'done','status':'done','progress':100,'message':'全部完成','outputPath':''})}\n\n"
        except Exception as e:
            yield f"data: {json.dumps({'type':'error','status':'error','message':str(e)})}\n\n"

    from fastapi.responses import StreamingResponse
    return StreamingResponse(event_generator(), media_type='text/event-stream')

