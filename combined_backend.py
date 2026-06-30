"""
案件处理启动器 - 合并后端服务 (Render.com 云部署版本)
包含附件处理器的完整 Node.js API 兼容层
"""
import os, sys, uuid, json, io, re, base64, urllib.request, asyncio
from datetime import datetime
from pathlib import Path
import tempfile
import zipfile

try:
    from dotenv import load_dotenv
    load_dotenv()
except ImportError:
    pass

BASE_DIR = Path(__file__).parent.resolve()
sys.path.insert(0, str(BASE_DIR))

from modules.config import HAS_MODULES
from modules.app_core import app
import modules.routes
from modules.routes import _generated_reports
from modules.template_data import PRESET_TEMPLATES, _get_preset_template

from fastapi.responses import JSONResponse, FileResponse, HTMLResponse
from fastapi import HTTPException, Request, UploadFile, File, Form
from fastapi.middleware.cors import CORSMiddleware

TEMP_DIR = Path(tempfile.gettempdir()) / 'case_processor'
TEMP_DIR.mkdir(exist_ok=True)
UPLOAD_DIR = TEMP_DIR / 'uploads'
OUTPUT_DIR = TEMP_DIR / 'output'
UPLOAD_DIR.mkdir(exist_ok=True)
OUTPUT_DIR.mkdir(exist_ok=True)

MIME_TYPES = {
    '.html': 'text/html; charset=utf-8', '.css': 'text/css', '.js': 'application/javascript',
    '.json': 'application/json', '.png': 'image/png', '.jpg': 'image/jpeg',
    '.jpeg': 'image/jpeg', '.gif': 'image/gif', '.svg': 'image/svg+xml', '.ico': 'image/x-icon',
    '.docx': 'application/vnd.openxmlformats-officedocument.wordprocessingml.document',
}

# ============ 附件处理器状态 ============
_attachment_images = []
_attachment_current_ranges = []
_sse_clients = set()

# 默认附件名称（与workflow.js一致）
DEFAULT_ATTACHMENT_NAMES = [
    "案卷资料记载（报案记录）及保单",
    "惠州市中拓模塑科技有限公司车间现场照片",
    "大门监控视频截图",
    "刘华妹询问笔录",
    "马少康个体诊所走访记录",
    "走访东莞市松山湖中心医院",
    "走访社保局",
    "黄昭伯身份证",
    "受益人身份证",
    "认定工伤决定书",
    "工亡待遇核定单",
    "死亡证明",
]

DEFAULT_RANGES = [
    {"start": 0, "end": 2, "idx": 0},
    {"start": 3, "end": 18, "idx": 1},
    {"start": 19, "end": 20, "idx": 2},
    {"start": 21, "end": 24, "idx": 3},
    {"start": 25, "end": 25, "idx": 4},
    {"start": 26, "end": 46, "idx": 5},
    {"start": 47, "end": 48, "idx": 6},
    {"start": 49, "end": 49, "idx": 7},
    {"start": 50, "end": 50, "idx": 8},
    {"start": 51, "end": 51, "idx": 9},
    {"start": 52, "end": 52, "idx": 10},
    {"start": 53, "end": 53, "idx": 11},
]


# ============ SSE 进度推送 ============
def sse_broadcast(data):
    msg = f"data: {json.dumps(data)}\n\n"
    dead = set()
    for client in _sse_clients:
        try:
            client.write(msg.encode())
        except:
            dead.add(client)
    _sse_clients -= dead


@app.get('/api/progress')
async def sse_progress(request: Request):
    async def event_generator():
        client_id = id(request)
        queue = asyncio.Queue()
        _sse_clients.add(client_id)
        try:
            # Send initial connected message
            yield f"data: {json.dumps({'type': 'connected', 'message': 'SSE连接已建立'})}\n\n"
            while True:
                try:
                    data = await asyncio.wait_for(queue.get(), timeout=30)
                    yield f"data: {json.dumps(data)}\n\n"
                except asyncio.TimeoutError:
                    yield ": keepalive\n\n"
        except:
            pass
        finally:
            _sse_clients.discard(client_id)

    from fastapi.responses import StreamingResponse
    return StreamingResponse(event_generator(), media_type="text/event-stream",
                             headers={"Cache-Control": "no-cache", "Connection": "keep-alive"})


# ============ 首页 ============
@app.get('/')
async def serve_index():
    for p in [BASE_DIR / 'index.html', BASE_DIR / 'templates' / 'index.html']:
        if p.exists():
            return HTMLResponse(p.read_text(encoding='utf-8'))
    return HTMLResponse('<h1>案件处理启动器</h1>')


# ============ 静态文件 ============
@app.get('/static/{file_path:path}')
async def serve_static(file_path: str):
    static_dir = BASE_DIR / 'static'
    full_path = (static_dir / file_path).resolve()
    if not str(full_path).startswith(str(static_dir.resolve())) or not full_path.is_file():
        return HTMLResponse('', status_code=404)
    return FileResponse(str(full_path), media_type=MIME_TYPES.get(full_path.suffix.lower(), 'application/octet-stream'))


@app.get('/tools/{file_path:path}')
async def serve_tools(file_path: str):
    tools_dir = BASE_DIR / 'tools'
    full_path = (tools_dir / file_path).resolve()
    if not str(full_path).startswith(str(tools_dir.resolve())):
        return HTMLResponse('Forbidden', status_code=403)
    if full_path.is_file():
        return FileResponse(str(full_path), media_type=MIME_TYPES.get(full_path.suffix.lower(), 'application/octet-stream'))
    if '.' not in file_path:
        html_path = full_path.with_suffix('.html')
        if html_path.is_file():
            return FileResponse(str(html_path), media_type=MIME_TYPES.get(html_path.suffix.lower(), 'application/octet-stream'))
    return HTMLResponse('Not Found', status_code=404)


@app.get('/temp/{filename}')
async def serve_temp(filename: str):
    """提供临时目录的文件（附件处理器缩略图用）"""
    filepath = TEMP_DIR / filename
    if filepath.is_file():
        return FileResponse(str(filepath), media_type='image/jpeg' if filepath.suffix.lower() in ['.jpg','.jpeg'] else 'image/png')
    filepath2 = UPLOAD_DIR / filename
    if filepath2.is_file():
        return FileResponse(str(filepath2), media_type='image/jpeg')
    raise HTTPException(404, 'File not found')


@app.get('/report-standard')
async def serve_report_standard():
    tpl_path = BASE_DIR / 'templates' / 'index.html'
    if tpl_path.exists():
        return HTMLResponse(tpl_path.read_text(encoding='utf-8'))
    return HTMLResponse('报告标准化页面', status_code=200)


# ============ 附件处理器 API ============

@app.get('/api/health')
async def health_check():
    return {'status': 'ok', 'service': 'combined-backend', 'version': '2.0.0'}


@app.get('/api/default-names')
async def get_default_names():
    """返回默认附件名称（附件处理器用）"""
    return {"names": DEFAULT_ATTACHMENT_NAMES, "ranges": DEFAULT_RANGES}


@app.get('/api/ranges')
async def get_ranges():
    """获取当前映射规则"""
    global _attachment_current_ranges
    if not _attachment_current_ranges:
        _attachment_current_ranges = [dict(r) for r in DEFAULT_RANGES]
    return {"ranges": _attachment_current_ranges}


@app.put('/api/ranges')
async def update_ranges(request: Request):
    """更新映射规则"""
    global _attachment_current_ranges
    data = await request.json()
    if not data or 'ranges' not in data:
        raise HTTPException(400, '请提供映射规则')
    _attachment_current_ranges = data['ranges']
    return {"success": True, "ranges": _attachment_current_ranges}


@app.post('/api/ranges/reset')
async def reset_ranges():
    """重置映射规则为默认值"""
    global _attachment_current_ranges
    _attachment_current_ranges = [dict(r) for r in DEFAULT_RANGES]
    return {"success": True, "ranges": _attachment_current_ranges}


@app.post('/api/upload-images')
async def upload_images(request: Request):
    """上传图片（附件处理器用）"""
    global _attachment_images
    try:
        form = await request.form()
        images = []
        for key in form.keys():
            value = form.get(key)
            if value and hasattr(value, 'read'):
                content = await value.read()
                if content:
                    ext = Path(getattr(value, 'filename', 'image.jpg')).suffix or '.jpg'
                    filename = f"img_{uuid.uuid4().hex}{ext}"
                    filepath = UPLOAD_DIR / filename
                    with open(filepath, 'wb') as f:
                        f.write(content)
                    images.append({
                        "index": len(_attachment_images) + len(images),
                        "name": getattr(value, 'filename', filename),
                        "size": len(content),
                        "filepath": str(filepath)
                    })
        
        # Also check for 'files' field (batched uploads)
        try:
            files = form.getlist('files')
            for f in files:
                if hasattr(f, 'read'):
                    content = await f.read()
                    if content:
                        ext = Path(getattr(f, 'filename', 'image.jpg')).suffix or '.jpg'
                        filename = f"img_{uuid.uuid4().hex}{ext}"
                        filepath = UPLOAD_DIR / filename
                        with open(filepath, 'wb') as f2:
                            f2.write(content)
                        images.append({
                            "index": len(_attachment_images) + len(images),
                            "name": getattr(f, 'filename', filename),
                            "size": len(content),
                            "filepath": str(filepath)
                        })
        except:
            pass

        _attachment_images.extend(images)
        sse_broadcast({"type": "images", "status": "uploaded", "count": len(images), "message": f"已上传 {len(images)} 张图片"})
        
        return {"success": True, "count": len(images), "images": images}
    except Exception as e:
        raise HTTPException(500, str(e))


@app.post('/api/upload-docx')
async def upload_docx(request: Request):
    """上传DOCX模板（附件处理器用）- 提取文字和图片尺寸"""
    try:
        form = await request.form()
        file = None
        for key in form.keys():
            value = form.get(key)
            if value and hasattr(value, 'read'):
                file = value
                break
        if not file:
            raise HTTPException(400, '请上传DOCX文件')
        
        content = await file.read()
        
        # 提取文字
        from docx import Document
        doc = Document(io.BytesIO(content))
        text_content = '\n'.join([p.text for p in doc.paragraphs if p.text.strip()])
        
        # 提取图片尺寸
        image_sizes = []
        try:
            z = zipfile.ZipFile(io.BytesIO(content))
            for name in z.namelist():
                if name.startswith('word/media/'):
                    img_data = z.read(name)
                    from PIL import Image as PILImage
                    img = PILImage.open(io.BytesIO(img_data))
                    image_sizes.append({
                        "name": name.replace('word/media/', ''),
                        "width": img.width,
                        "height": img.height,
                        "size": len(img_data),
                        "orientation": "portrait" if img.height > img.width else "landscape"
                    })
        except:
            pass
        
        # 提取标题
        title_lines = [l for l in text_content.split('\n') if len(l.strip()) > 5]
        title = title_lines[0] if title_lines else '未命名文档'
        
        return {
            "success": True,
            "content": text_content,
            "title": title,
            "header": {"text": title, "hasIcon": False},
            "imageSizes": image_sizes
        }
    except Exception as e:
        raise HTTPException(500, str(e))


@app.post('/api/upload-folder')
async def upload_folder(request: Request):
    """上传文件夹（自动识别附件映射）"""
    try:
        form = await request.form()
        folder_mapping = json.loads(form.get('folderMapping', '[]'))
        files = []
        try:
            for f in form.getlist('files'):
                if hasattr(f, 'read'):
                    content = await f.read()
                    if content:
                        ext = Path(getattr(f, 'filename', 'image.jpg')).suffix or '.jpg'
                        filename = f"img_{uuid.uuid4().hex}{ext}"
                        filepath = UPLOAD_DIR / filename
                        with open(filepath, 'wb') as f2:
                            f2.write(content)
                        files.append({
                            "index": len(_attachment_images) + len(files),
                            "name": getattr(f, 'filename', filename),
                            "size": len(content),
                            "filepath": str(filepath)
                        })
        except:
            pass
        
        _attachment_images.extend(files)
        
        # 生成自动映射
        auto_ranges = []
        for folder in folder_mapping:
            idx = 0
            folder_name = folder.get('folderName', '')
            import re as _re
            m = _re.search(r'附件([一二三四五六七八九十\d]+)', folder_name)
            if m:
                num_str = m.group(1)
                chn = {'一':0,'二':1,'三':2,'四':3,'五':4,'六':5,'七':6,'八':7,'九':8,'十':9}
                idx = chn.get(num_str, int(num_str)-1 if num_str.isdigit() else 0)
            indexes = folder.get('fileIndexes', [])
            if indexes:
                auto_ranges.append({"start": min(indexes), "end": max(indexes), "idx": idx})
        
        return {
            "success": True,
            "folders": len(folder_mapping),
            "totalImages": len(files),
            "images": files,
            "ranges": auto_ranges,
            "hasAutoMapping": len(auto_ranges) > 0
        }
    except Exception as e:
        raise HTTPException(500, str(e))


@app.post('/api/generate-markdown')
async def generate_markdown(request: Request):
    """生成Markdown格式的附件报告"""
    try:
        data = await request.json()
        doc_content = data.get('docContent', '')
        attachment_names = data.get('attachmentNames', '')
        custom_ranges = data.get('customRanges')
        images_data = data.get('images', [])
        
        # 解析附件名称
        names = parse_attachment_names(attachment_names)
        
        # 构建附件
        ranges = custom_ranges if custom_ranges else DEFAULT_RANGES
        attachments = build_attachments(images_data, names, ranges)
        
        # 生成Markdown
        md = generate_markdown_content(attachments)
        
        sse_broadcast({"type": "markdown", "status": "done", "message": "Markdown生成完成"})
        
        return {"success": True, "markdown": md, "attachments": attachments, "header": {"text": "", "hasIcon": False}}
    except Exception as e:
        raise HTTPException(500, str(e))


@app.post('/api/convert-word')
async def convert_word(request: Request):
    """将Markdown转换为Word文档"""
    try:
        data = await request.json()
        markdown = data.get('markdown', '')
        images_data = data.get('images', [])
        
        from docx import Document as DocxDocument
        from docx.shared import Inches, Pt
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        
        doc = DocxDocument()
        section = doc.sections[0]
        
        # 解析Markdown中的图片引用
        img_tags = re.findall(r'<img[^>]+src="([^"]+)"', markdown)
        
        # 对于每个图片标签
        for img_url in img_tags:
            # 如果是本地文件路径
            for img in images_data:
                fp = img.get('filepath', '')
                if fp and os.path.exists(fp):
                    try:
                        doc.add_picture(fp, width=Inches(5.5))
                        last_p = doc.paragraphs[-1]
                        last_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    except:
                        pass
        
        output = io.BytesIO()
        doc.save(output)
        output.seek(0)
        
        output_filename = f"output_{uuid.uuid4().hex}.docx"
        output_path = TEMP_DIR / output_filename
        with open(output_path, 'wb') as f:
            f.write(output.read())
        output.seek(0)
        
        sse_broadcast({"type": "word", "status": "done", "message": "Word文档生成完成！", "outputPath": output_filename})
        
        return {"success": True, "outputPath": str(output_path), "size": os.path.getsize(output_path)}
    except Exception as e:
        raise HTTPException(500, str(e))


@app.post('/api/one-click')
async def one_click(request: Request):
    """一键全流程：上传→映射→Markdown→Word"""
    try:
        data = await request.json()
        images_data = data.get('images', [])
        attachment_names = data.get('attachmentNames', '')
        custom_ranges = data.get('customRanges')
        header = data.get('header', {})
        docx_image_sizes = data.get('docxImageSizes', [])
        
        total = len(images_data)
        
        sse_broadcast({"type": "oneclick", "phase": "build", "status": "processing", "progress": 0, "message": "构建附件映射..."})
        await asyncio.sleep(0.1)
        
        # 构建附件
        names = parse_attachment_names(attachment_names)
        ranges = custom_ranges if custom_ranges else DEFAULT_RANGES
        attachments = build_attachments(images_data, names, ranges)
        
        sse_broadcast({"type": "oneclick", "phase": "markdown", "status": "processing", "progress": 15, "message": "生成Markdown..."})
        await asyncio.sleep(0.1)
        
        # 生成Markdown
        md = generate_markdown_content(attachments)
        
        sse_broadcast({"type": "oneclick", "phase": "images", "status": "processing", "progress": 45, "message": f"处理 {total} 张图片..."})
        
        # 处理图片
        image_results = []
        for i, img in enumerate(images_data):
            fp = img.get('filepath', '')
            result = {"success": False, "filepath": None}
            if fp and os.path.exists(fp):
                result = {"success": True, "filepath": fp}
            image_results.append(result)
            progress = 45 + int(((i+1) / total) * 30)
            sse_broadcast({"type": "oneclick", "phase": "images", "status": "progress", "current": i+1, "total": total, "progress": progress})
        
        sse_broadcast({"type": "oneclick", "phase": "word", "status": "processing", "progress": 78, "message": "生成Word文档..."})
        
        # 生成Word
        from docx import Document as DocxDocument
        from docx.shared import Inches, Pt
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from PIL import Image as PILImage
        
        doc = DocxDocument()
        
        for i, att in enumerate(attachments):
            # 附件编号
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            r = p.add_run(f"附件{['一','二','三','四','五','六','七','八','九','十','十一','十二'][i] if i < 12 else str(i+1)}")
            r.bold = True; r.font.size = Pt(11)
            
            # 附件名称
            p2 = doc.add_paragraph()
            p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r2 = p2.add_run(att['name'])
            r2.bold = True; r2.font.size = Pt(12)
            
            # 图片
            for img in att.get('images', []):
                fp = img.get('filepath', '')
                if fp and os.path.exists(fp):
                    try:
                        pil = PILImage.open(fp)
                        w, h = pil.size
                        if w > h:
                            doc.add_picture(fp, width=Inches(7.38))
                        else:
                            doc.add_picture(fp, width=Inches(5.53))
                        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
                    except:
                        pass
        
        output_filename = f"output_{uuid.uuid4().hex}.docx"
        output_path = TEMP_DIR / output_filename
        doc.save(str(output_path))
        
        sse_broadcast({
            "type": "oneclick", "phase": "done", "status": "success",
            "progress": 100, "message": "全流程完成！Word文档已生成",
            "outputPath": output_filename,
            "stats": {"images": total, "attachments": len(attachments), "size": os.path.getsize(output_path)}
        })
        
        return {"success": True, "markdown": md, "attachments": attachments, "outputPath": str(output_path), "filename": output_filename, "size": os.path.getsize(output_path)}
    except Exception as e:
        sse_broadcast({"type": "oneclick", "phase": "error", "status": "error", "message": f"全流程失败：{str(e)}"})
        raise HTTPException(500, f'全流程失败: {str(e)}')


@app.get('/api/download/{filename}')
async def download_file(filename: str):
    """下载附件处理器生成的文件"""
    filepath = TEMP_DIR / filename
    if not filepath.is_file():
        filepath = OUTPUT_DIR / filename
    if not filepath.is_file():
        filepath = UPLOAD_DIR / filename
    if not filepath.is_file():
        raise HTTPException(404, '文件不存在')
    return FileResponse(str(filepath), filename=filename, media_type=MIME_TYPES.get(filepath.suffix.lower(), 'application/octet-stream'))


@app.post('/api/cleanup')
async def cleanup_temp():
    """清理临时文件"""
    cleaned = 0
    now = datetime.now().timestamp()
    for d in [TEMP_DIR, UPLOAD_DIR]:
        for f in d.glob('*'):
            if f.is_file() and (now - f.stat().st_mtime) > 3600:
                try:
                    f.unlink()
                    cleaned += 1
                except:
                    pass
    return {"success": True, "cleaned": cleaned}


# ============ 附件处理器辅助函数 ============

def parse_attachment_names(names_text):
    """解析附件名称"""
    if not names_text or not names_text.strip():
        return DEFAULT_ATTACHMENT_NAMES[:]
    lines = [l.strip() for l in names_text.split('\n') if l.strip()]
    result = DEFAULT_ATTACHMENT_NAMES[:]
    for i, line in enumerate(lines):
        if i >= len(result):
            break
        m = re.match(r'^附件[一二三四五六七八九十\d]+[：:.]?\s*(.*)', line)
        if m and m.group(1):
            result[i] = m.group(1)
        elif not line.startswith('附件'):
            result[i] = line
    return result


def build_attachments(images, names, ranges):
    """构建附件数据结构"""
    result = []
    chn = ['一','二','三','四','五','六','七','八','九','十','十一','十二']
    for i, r in enumerate(ranges):
        img_list = []
        for j in range(r['start'], r['end']+1):
            if j < len(images):
                img_list.append(images[j])
        name = names[r['idx']] if r['idx'] < len(names) else f"附件{chn[r['idx']] if r['idx'] < 12 else str(r['idx']+1)}"
        result.append({
            "index": i+1,
            "indexChinese": chn[i] if i < 12 else str(i+1),
            "name": name,
            "range": f"{r['start']}-{r['end']}",
            "count": len(img_list),
            "images": img_list,
            "urls": [img.get('filepath', '') for img in img_list]
        })
    return result


def generate_markdown_content(attachments):
    """生成Markdown格式的附件报告"""
    lines = []
    for att in attachments:
        lines.append(f'<span style="float:right;color:red;font-weight:bold;">附件{att["indexChinese"]}</span>')
        lines.append('')
        lines.append(f'### <span style="text-align:center;display:block;">{att["name"]}</span>')
        lines.append('')
        for img in att.get('images', []):
            fp = img.get('filepath', '')
            if fp and os.path.exists(fp):
                lines.append(f'<img src="{fp}" width="100%" />')
                lines.append('')
    return '\n'.join(lines)


@app.post('/api/test')
async def test_endpoint():
    return {'success': True, 'message': '后端服务正常运行'}


# ============ 服务器启动 ============
if __name__ == '__main__':
    import uvicorn
    port = int(os.environ.get('PORT', 10000))
    host = os.environ.get('HOST', '0.0.0.0')
    print(f'=== 后端服务启动在 {host}:{port} ===')
    uvicorn.run(app, host=host, port=port, log_level='info', access_log=False, timeout_keep_alive=600)