#!/usr/bin/env python3
"""简单检查生成的报告"""

import sys
from pathlib import Path
from docx import Document

output_dir = Path(__file__).parent / "output"
test_file = output_dir / "test_report_generation.docx"

if not test_file.exists():
    print("Test file not found")
    sys.exit(1)

print("Found test file")
print(f"File size: {test_file.stat().st_size} bytes")

doc = Document(str(test_file))
print(f"\nParagraphs: {len(doc.paragraphs)}")
print(f"Tables: {len(doc.tables)}")

print("\nDocument content:")
for i, p in enumerate(doc.paragraphs):
    if p.text.strip():
        print(f"{i+1}. {p.text}")
