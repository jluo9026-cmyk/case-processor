# -*- coding: utf-8 -*-
import sys
sys.stdout.reconfigure(encoding='utf-8')
sys.stderr.reconfigure(encoding='utf-8')

from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn

doc = Document('e:/桌面/测试/深圳人保-王玉香报告.docx')

with open('e:/案件处理启动器/analyze_output.txt', 'w', encoding='utf-8') as f:
    f.write('=== 文档样式分析 ===\n')
    f.write(f'总段落数: {len(doc.paragraphs)}\n\n')
    
    for i, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        if not text:
            continue
        
        style_name = para.style.name if para.style else 'None'
        alignment = str(para.alignment) if para.alignment else 'None'
        
        if para.runs:
            run = para.runs[0]
            font_name = run.font.name
            font_size = run.font.size
            font_size_pt = font_size.pt if font_size else 'inherit'
            bold = run.font.bold
            
            east_asian = 'N/A'
            try:
                if run._element.rPr is not None and run._element.rPr.rFonts is not None:
                    east_asian = run._element.rPr.rFonts.get(qn('w:eastAsia'))
            except:
                pass
        else:
            font_name = 'N/A'
            font_size_pt = 'N/A'
            bold = 'N/A'
            east_asian = 'N/A'
        
        preview = text[:60]
        f.write(f'P{i:3d}: 样式={style_name:20s} 对齐={alignment:15s} 字体={str(font_name):15s} 东亚={str(east_asian):15s} 字号={str(font_size_pt):8s} 加粗={str(bold):5s}\n')
        f.write(f'      内容: {preview}\n\n')

print('Done!')
