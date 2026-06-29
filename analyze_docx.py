# -*- coding: utf-8 -*-
from docx import Document
from docx.shared import Pt, Inches, Cm
from docx.oxml.ns import qn

doc = Document(r'e:\桌面\测试\深圳人保-王玉香报告.docx')

print('=== 文档样式分析 ===')
print(f'总段落数: {len(doc.paragraphs)}')
print()

for i, para in enumerate(doc.paragraphs):
    text = para.text.strip()
    if not text:
        continue
    
    style_name = para.style.name if para.style else 'None'
    alignment = str(para.alignment) if para.alignment else 'None'
    
    # 获取第一个run的字体信息
    if para.runs:
        run = para.runs[0]
        font_name = run.font.name
        font_size = run.font.size
        if font_size:
            font_size_pt = font_size.pt
        else:
            font_size_pt = 'inherit'
        bold = run.font.bold
        
        # 获取东亚字体
        east_asian = 'N/A'
        try:
            if run._element.rPr is not None and run._element.rPr.rFonts is not None:
                east_asian = run._element.rPr.rFonts.get(qn('w:eastAsia'))
        except:
            pass
    else:
        font_name = 'N/A'
        font_size_pt = 'N/A'
        bold = 'N/A'
        east_asian = 'N/A'
    
    preview = text[:60]
    print(f'P{i:3d}: 样式={style_name:20s} 对齐={alignment:15s} 字体={str(font_name):15s} 东亚={str(east_asian):15s} 字号={str(font_size_pt):8s} 加粗={str(bold):5s}')
    print(f'      内容: {preview}')
    print()
